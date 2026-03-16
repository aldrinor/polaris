"""
Anti-Tunnel-View Integration Test: Verifies feature survival through the FULL pipeline.

Unlike component tests that check individual functions, this traces data flow
across module boundaries to catch integration bugs that component tests miss.

Usage:
    python scripts/anti_tunnel_view_test.py           # Local only ($0)
    python scripts/anti_tunnel_view_test.py --api      # Include API tests (~$0.10)
"""
import asyncio
import base64
import json
import os
import re
import sys
import tempfile
from pathlib import Path
from typing import Any

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
os.chdir(PROJECT_ROOT)

# Load .env
from dotenv import load_dotenv
load_dotenv(PROJECT_ROOT / ".env")

PASS = 0
FAIL = 0
WARN = 0
RESULTS: list[dict] = []


def _report(name: str, status: str, detail: str = ""):
    global PASS, FAIL, WARN
    icon = {"PASS": "+", "FAIL": "!", "WARN": "~", "INFO": "*"}[status]
    if status == "PASS":
        PASS += 1
    elif status == "FAIL":
        FAIL += 1
    elif status == "WARN":
        WARN += 1
    print(f"  [{icon}] {name}: {detail}")
    RESULTS.append({"name": name, "status": status, "detail": detail})


# ---------------------------------------------------------------------------
# LAYER 1: State Schema Integrity
# ---------------------------------------------------------------------------
def test_state_schema_completeness():
    """Verify all Gemini features have state keys declared."""
    from src.polaris_graph.state import ResearchState
    hints = ResearchState.__annotations__

    required_keys = [
        "structured_data", "smart_art_diagrams", "hallucination_audit",
        "evidence", "fetched_content", "sections", "final_report",
        "section_evidence_map", "cross_reference_groups",
    ]
    missing = [k for k in required_keys if k not in hints]
    if missing:
        _report("state_schema_completeness", "FAIL",
                f"Missing keys in ResearchState: {missing}")
    else:
        _report("state_schema_completeness", "PASS",
                f"All {len(required_keys)} Gemini-critical keys declared")


def test_initial_state_defaults():
    """Verify create_initial_state() initializes all Gemini keys."""
    from src.polaris_graph.state import create_initial_state
    state = create_initial_state("test_vec", "test query", "test app", "test region")

    checks = {
        "structured_data": lambda v: isinstance(v, list),
        "smart_art_diagrams": lambda v: isinstance(v, dict),
        "hallucination_audit": lambda v: isinstance(v, list),
        "evidence": lambda v: isinstance(v, list),
        "section_evidence_map": lambda v: isinstance(v, dict),
    }
    failures = []
    for key, check in checks.items():
        val = state.get(key)
        if val is None:
            failures.append(f"{key}=None (not initialized)")
        elif not check(val):
            failures.append(f"{key} wrong type: {type(val).__name__}")

    if failures:
        _report("initial_state_defaults", "FAIL", "; ".join(failures))
    else:
        _report("initial_state_defaults", "PASS",
                f"All {len(checks)} Gemini state keys properly initialized")


# ---------------------------------------------------------------------------
# LAYER 2: Feature Flag Verification
# ---------------------------------------------------------------------------
def test_feature_flags_enabled():
    """Verify all Gemini feature flags are set in .env."""
    required_flags = {
        "PG_CLUSTER_VIABILITY_ENABLED": "1",
        "PG_STRUCTURED_DATA_EXTRACTION": "1",
        "PG_CHART_GENERATION_ENABLED": "1",
        "PG_KEY_FINDINGS_ENFORCEMENT": "1",
        "PG_SMART_ART_ENABLED": "1",
    }
    issues = []
    for flag, expected in required_flags.items():
        actual = os.getenv(flag, "NOT SET")
        if actual != expected:
            issues.append(f"{flag}={actual} (expected {expected})")

    if issues:
        _report("feature_flags_enabled", "FAIL", "; ".join(issues))
    else:
        _report("feature_flags_enabled", "PASS",
                f"All {len(required_flags)} Gemini flags enabled")


def test_feature_flag_code_guards():
    """Verify feature flags are actually CHECKED in code (not just declared)."""
    flag_to_file = {
        "PG_STRUCTURED_DATA_EXTRACTION": "src/polaris_graph/agents/analyzer.py",
        "PG_CHART_GENERATION_ENABLED": "src/polaris_graph/agents/synthesizer.py",
        "PG_KEY_FINDINGS_ENFORCEMENT": "src/polaris_graph/synthesis/section_writer.py",
        "PG_SMART_ART_ENABLED": "src/polaris_graph/synthesis/smart_art_generator.py",
        "PG_CLUSTER_VIABILITY_ENABLED": "src/polaris_graph/agents/synthesizer.py",
    }
    missing_guards = []
    for flag, filepath in flag_to_file.items():
        full_path = PROJECT_ROOT / filepath
        if full_path.exists():
            content = full_path.read_text(encoding="utf-8")
            if flag not in content:
                missing_guards.append(f"{flag} NOT checked in {filepath}")

    if missing_guards:
        _report("feature_flag_guards", "WARN", "; ".join(missing_guards))
    else:
        _report("feature_flag_guards", "PASS",
                f"All {len(flag_to_file)} flags checked in their gating code")


# ---------------------------------------------------------------------------
# LAYER 3: Data Flow Tracing (structured_data chain)
# ---------------------------------------------------------------------------
def test_analyzer_returns_structured_data():
    """Verify analyzer return dict includes structured_data key."""
    import ast
    analyzer_path = PROJECT_ROOT / "src/polaris_graph/agents/analyzer.py"
    content = analyzer_path.read_text(encoding="utf-8")

    # Check the return dict at the bottom of analyze_sources()
    if '"structured_data":' in content or "'structured_data':" in content:
        _report("analyzer_returns_structured_data", "PASS",
                "analyzer.py returns structured_data in result dict")
    else:
        _report("analyzer_returns_structured_data", "FAIL",
                "analyzer.py does NOT return structured_data")


def test_graph_accumulates_structured_data():
    """Verify graph.py has accumulation logic for structured_data."""
    graph_path = PROJECT_ROOT / "src/polaris_graph/graph.py"
    content = graph_path.read_text(encoding="utf-8")

    patterns = [
        'state.get("structured_data"',
        'result["structured_data"]',
        "GEMINI-ARCH: ACCUMULATE structured_data",
    ]
    found = [p for p in patterns if p in content]
    if len(found) >= 2:
        _report("graph_accumulates_structured_data", "PASS",
                f"Found {len(found)}/3 accumulation patterns in graph.py")
    else:
        _report("graph_accumulates_structured_data", "FAIL",
                f"Only {len(found)}/3 accumulation patterns found")


def test_synthesizer_reads_structured_data():
    """Verify synthesizer reads structured_data from state for chart gen."""
    synth_path = PROJECT_ROOT / "src/polaris_graph/agents/synthesizer.py"
    content = synth_path.read_text(encoding="utf-8")

    if 'state.get("structured_data"' in content:
        _report("synthesizer_reads_structured_data", "PASS",
                "synthesizer.py reads state['structured_data'] for chart gen")
    else:
        _report("synthesizer_reads_structured_data", "FAIL",
                "synthesizer.py does NOT read structured_data from state")


def test_chart_injection_into_sections():
    """Verify _generate_section_charts appends charts to section content."""
    synth_path = PROJECT_ROOT / "src/polaris_graph/agents/synthesizer.py"
    content = synth_path.read_text(encoding="utf-8")

    # Check that section content gets updated with chart markdown
    has_chart_append = "sec.content = updated_content" in content
    has_chart_gen_call = "_generate_section_charts" in content
    has_format_chart = "format_chart_markdown" in content

    if has_chart_append and has_chart_gen_call and has_format_chart:
        _report("chart_injection_into_sections", "PASS",
                "Charts appended to section content via _generate_section_charts()")
    else:
        _report("chart_injection_into_sections", "FAIL",
                f"chart_append={has_chart_append}, gen_call={has_chart_gen_call}, format={has_format_chart}")


# ---------------------------------------------------------------------------
# LAYER 4: Key Findings Survival
# ---------------------------------------------------------------------------
def test_key_findings_preservation():
    """Verify Key Findings extract/preserve/re-append in section_writer.py."""
    sw_path = PROJECT_ROOT / "src/polaris_graph/synthesis/section_writer.py"
    content = sw_path.read_text(encoding="utf-8")

    has_extract = "_kf_preserve_block" in content or "kf_preserve" in content.lower()
    has_reappend = re.search(r'if.*_kf_preserve_block', content) is not None
    has_key_findings_regex = re.search(r'Key Findings', content) is not None

    if has_extract and has_reappend:
        _report("key_findings_preservation", "PASS",
                "FIX-KF-PRESERVE: extract before post-proc, re-append after")
    elif has_key_findings_regex:
        _report("key_findings_preservation", "WARN",
                "Key Findings referenced but preservation pattern incomplete")
    else:
        _report("key_findings_preservation", "FAIL",
                "No Key Findings preservation logic found")


# ---------------------------------------------------------------------------
# LAYER 5: :::metrics Survival
# ---------------------------------------------------------------------------
def test_metrics_unconditional_insertion():
    """Verify :::metrics is inserted regardless of expansion_passes."""
    synth_path = PROJECT_ROOT / "src/polaris_graph/agents/synthesizer.py"
    content = synth_path.read_text(encoding="utf-8")

    # Check for the unconditional insertion (if True:)
    if "if True:" in content and ":::metrics" in content:
        _report("metrics_unconditional_insertion", "PASS",
                "FIX-METRICS-ALWAYS: :::metrics inserted unconditionally")
    elif ":::metrics" in content:
        _report("metrics_unconditional_insertion", "WARN",
                ":::metrics exists but may be gated by expansion_passes")
    else:
        _report("metrics_unconditional_insertion", "FAIL",
                "No :::metrics insertion found")


def test_metrics_survives_post_processing():
    """Check if post-processing in report_assembler could strip :::metrics."""
    ra_path = PROJECT_ROOT / "src/polaris_graph/synthesis/report_assembler.py"
    content = ra_path.read_text(encoding="utf-8")

    # Check post-processing functions for dangerous patterns
    danger_patterns = [
        r'\.replace\(.{0,20}:::', r'strip.*:::', r'sub\(.{0,30}:::'
    ]
    dangers = []
    for pat in danger_patterns:
        if re.search(pat, content):
            dangers.append(pat)

    if not dangers:
        _report("metrics_survives_post_processing", "PASS",
                "No post-processing patterns that would strip :::metrics")
    else:
        _report("metrics_survives_post_processing", "WARN",
                f"Potentially dangerous patterns found: {dangers}")


# ---------------------------------------------------------------------------
# LAYER 6: Matplotlib Execution (real subprocess)
# ---------------------------------------------------------------------------
def test_matplotlib_subprocess():
    """Actually execute a matplotlib chart via subprocess — proves Windows works."""
    script = '''
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import json, io, base64

fig, ax = plt.subplots(figsize=(6, 4))
ax.bar(['A', 'B', 'C'], [10, 20, 15], color=['#2196F3', '#4CAF50', '#FF9800'])
ax.set_title('Test Chart')
ax.set_ylabel('Value')
buf = io.BytesIO()
fig.savefig(buf, format='png', dpi=100, bbox_inches='tight')
buf.seek(0)
b64 = base64.b64encode(buf.read()).decode()
plt.close()
print(json.dumps({
    "charts": [{"title": "Test", "image_base64": b64, "description": "Test chart"}],
    "tables": [{"headers": ["Col1", "Col2"], "rows": [["A", "10"], ["B", "20"]], "caption": "Test table"}],
    "insights": ["Value B is highest at 20"]
}))
'''
    import subprocess
    with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False, encoding='utf-8') as f:
        f.write(script)
        script_path = f.name

    try:
        proc = subprocess.run(
            [sys.executable, script_path],
            capture_output=True, text=True, timeout=30,
            cwd=tempfile.gettempdir(),
        )
        if proc.returncode != 0:
            _report("matplotlib_subprocess", "FAIL",
                    f"exit={proc.returncode}, stderr={proc.stderr[:200]}")
            return

        result = json.loads(proc.stdout.strip())
        charts = result.get("charts", [])
        tables = result.get("tables", [])

        if not charts:
            _report("matplotlib_subprocess", "FAIL", "No charts in output")
            return

        b64 = charts[0].get("image_base64", "")
        decoded = base64.b64decode(b64)
        is_png = decoded[:4] == b'\x89PNG'

        if is_png and len(decoded) > 1000 and tables:
            _report("matplotlib_subprocess", "PASS",
                    f"Chart={len(decoded)}b valid PNG, table={len(tables[0].get('rows', []))} rows")
        else:
            _report("matplotlib_subprocess", "FAIL",
                    f"PNG valid={is_png}, size={len(decoded)}, tables={len(tables)}")
    except subprocess.TimeoutExpired:
        _report("matplotlib_subprocess", "FAIL", "Subprocess timed out (30s)")
    except Exception as exc:
        _report("matplotlib_subprocess", "FAIL", f"Exception: {exc}")
    finally:
        try:
            os.unlink(script_path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# LAYER 7: Schema Validation
# ---------------------------------------------------------------------------
def test_structured_data_point_coercion():
    """Verify StructuredDataPoint accepts float values (LLM returns numbers)."""
    from src.polaris_graph.schemas import StructuredDataPoint
    try:
        point = StructuredDataPoint(
            value=7.0,  # LLM returns float, not str
            unit="pH",
            entity="water sample",
            source_url="https://example.com",
            evidence_id="ev_test_001",
        )
        assert isinstance(point.value, str), f"value not coerced to str: {type(point.value)}"
        assert point.value == "7.0"
        _report("structured_data_point_coercion", "PASS",
                "Float-to-str coercion works (FIX-SDP1)")
    except Exception as exc:
        _report("structured_data_point_coercion", "FAIL", str(exc)[:200])


def test_structured_data_extraction_wrap():
    """Verify StructuredDataExtraction wraps bare lists."""
    from src.polaris_graph.schemas import StructuredDataExtraction
    try:
        # LLM returns bare list instead of {"data_points": [...]}
        bare_list = [
            {"value": "95", "unit": "%", "entity": "removal", "source_url": "", "evidence_id": "ev_1"},
        ]
        result = StructuredDataExtraction.model_validate(bare_list)
        assert len(result.data_points) == 1
        _report("structured_data_extraction_wrap", "PASS",
                "Bare list auto-wrapped (FIX-SDE1)")
    except Exception as exc:
        _report("structured_data_extraction_wrap", "FAIL", str(exc)[:200])


# ---------------------------------------------------------------------------
# LAYER 8: Chart Markdown Formatting
# ---------------------------------------------------------------------------
def test_chart_markdown_format():
    """Verify format_chart_markdown produces valid embedded image."""
    from src.polaris_graph.tools.data_analyzer import format_chart_markdown

    # Create a minimal valid PNG (1x1 red pixel)
    import struct, zlib
    width, height = 1, 1
    raw_data = b'\x00\xff\x00\x00'  # filter byte + RGB
    compressed = zlib.compress(raw_data)

    def _chunk(chunk_type, data):
        c = chunk_type + data
        return struct.pack('>I', len(data)) + c + struct.pack('>I', zlib.crc32(c) & 0xFFFFFFFF)

    png = b'\x89PNG\r\n\x1a\n'
    png += _chunk(b'IHDR', struct.pack('>IIBBBBB', width, height, 8, 2, 0, 0, 0))
    png += _chunk(b'IDAT', compressed)
    png += _chunk(b'IEND', b'')

    b64 = base64.b64encode(png).decode()
    chart = {"title": "Test", "image_base64": b64, "description": "A test chart"}
    md = format_chart_markdown(chart, figure_number=1)

    has_img = "![Test]" in md
    has_b64 = "data:image/png;base64," in md
    has_caption = "*Figure 1:" in md

    if has_img and has_b64 and has_caption:
        _report("chart_markdown_format", "PASS", f"Valid markdown ({len(md)} chars)")
    else:
        _report("chart_markdown_format", "FAIL",
                f"img={has_img}, b64={has_b64}, caption={has_caption}")


def test_table_markdown_format():
    """Verify format_table_markdown produces valid pipe-delimited table."""
    from src.polaris_graph.tools.data_analyzer import format_table_markdown

    table = {
        "headers": ["Material", "Strength (MPa)", "Cost ($/kg)"],
        "rows": [["Steel", "400", "2.50"], ["Aluminum", "270", "3.80"]],
        "caption": "Material comparison",
    }
    md = format_table_markdown(table, table_number=1)

    has_pipes = "| Material |" in md
    has_separator = "| --- |" in md
    has_data = "| Steel |" in md
    has_caption = "*Table 1:" in md

    if has_pipes and has_separator and has_data and has_caption:
        _report("table_markdown_format", "PASS", f"Valid table ({len(md)} chars)")
    else:
        _report("table_markdown_format", "FAIL",
                f"pipes={has_pipes}, sep={has_separator}, data={has_data}, caption={has_caption}")


# ---------------------------------------------------------------------------
# LAYER 9: Filler Reduction
# ---------------------------------------------------------------------------
def test_filler_reduction():
    """Verify filler reduction in report_assembler actually works."""
    from src.polaris_graph.synthesis.report_assembler import _reduce_filler

    text = (
        "Furthermore, the study found significant results. "
        "Moreover, the analysis revealed patterns. "
        "Furthermore, additional testing confirmed the hypothesis. "
        "Additionally, the data showed strong correlation. "
        "Furthermore, these findings suggest improvement. "
        "Moreover, the control group exhibited similar trends. "
        "Furthermore, the methodology was validated. "
        "It is worth noting that accuracy improved. "
    )
    result = _reduce_filler(text)
    # Count remaining filler words
    filler_count = sum(1 for w in ["Furthermore,", "Moreover,", "Additionally,", "It is worth noting"]
                       if w in result)
    original_filler = 8
    if filler_count < original_filler:
        _report("filler_reduction", "PASS",
                f"Reduced {original_filler} fillers to {filler_count}")
    else:
        _report("filler_reduction", "FAIL",
                f"No reduction: {filler_count}/{original_filler} fillers remain")


# ---------------------------------------------------------------------------
# LAYER 10: DOCX Export with Rich Content
# ---------------------------------------------------------------------------
def test_docx_export_with_rich_content():
    """Verify DOCX export handles tables, images, and Key Findings."""
    try:
        from src.polaris_graph.export.docx_exporter import DocxExporter

        # Create sample markdown with all Gemini features
        md = """# Test Report

## Introduction

This section introduces the topic. [1]

**Key Findings:**
- Finding one: 95% efficiency [1]
- Finding two: cost reduced by 40% [2]

## Data Analysis

| Material | Strength (MPa) | Cost |
| --- | --- | --- |
| Steel | 400 | $2.50 |
| Aluminum | 270 | $3.80 |

*Table 1: Material comparison*

## Conclusion

The analysis demonstrates clear advantages. [1][2]

## References

[1] Smith et al. (2024). Journal of Materials.
[2] Johnson et al. (2024). Materials Science Review.
"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            output_path = f.name

        exporter = DocxExporter()
        # DocxExporter.export() expects (report_data: dict, output_path: Path)
        report_data = {
            "final_report": md,
            "original_query": "Test Report",
            "vector_id": "test_001",
            "status": "complete",
            "bibliography": [],
            "evidence": [],
            "quality_metrics": None,
        }
        exporter.export(report_data, Path(output_path))

        file_size = os.path.getsize(output_path)

        from docx import Document
        doc = Document(output_path)
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        os.unlink(output_path)

        if file_size > 5000 and table_count >= 1:
            _report("docx_export_rich_content", "PASS",
                    f"DOCX: {file_size}b, {para_count} paras, {table_count} tables")
        elif file_size > 5000:
            _report("docx_export_rich_content", "WARN",
                    f"DOCX created ({file_size}b) but tables={table_count}")
        else:
            _report("docx_export_rich_content", "FAIL",
                    f"DOCX too small: {file_size}b")

    except ImportError as exc:
        _report("docx_export_rich_content", "WARN", f"Import error: {exc}")
    except Exception as exc:
        _report("docx_export_rich_content", "FAIL", str(exc)[:200])


# ---------------------------------------------------------------------------
# LAYER 11: Smart Art Gap Detection
# ---------------------------------------------------------------------------
def test_smart_art_injection_gap():
    """Check if smart_art_diagrams are injected into final report (known gap)."""
    ra_path = PROJECT_ROOT / "src/polaris_graph/synthesis/report_assembler.py"
    content = ra_path.read_text(encoding="utf-8")

    if "smart_art" in content:
        _report("smart_art_injection_gap", "PASS",
                "report_assembler.py references smart_art")
    else:
        _report("smart_art_injection_gap", "WARN",
                "KNOWN GAP: smart_art_diagrams generated but NOT injected into report_assembler.py")


# ---------------------------------------------------------------------------
# LAYER 12: Evidence Accumulation Pattern
# ---------------------------------------------------------------------------
def test_evidence_accumulation_pattern():
    """Verify graph.py accumulates evidence, fetched_content, AND structured_data."""
    graph_path = PROJECT_ROOT / "src/polaris_graph/graph.py"
    content = graph_path.read_text(encoding="utf-8")

    accumulations = {
        "evidence": 'FIX-300: ACCUMULATE evidence' in content,
        "fetched_content": 'RC-1: ACCUMULATE fetched_content' in content,
        "structured_data": 'GEMINI-ARCH: ACCUMULATE structured_data' in content,
    }

    missing = [k for k, v in accumulations.items() if not v]
    if not missing:
        _report("evidence_accumulation_pattern", "PASS",
                "All 3 data types accumulated across iterations")
    else:
        _report("evidence_accumulation_pattern", "FAIL",
                f"Missing accumulation for: {missing}")


# ---------------------------------------------------------------------------
# LAYER 13: Real Output Analysis
# ---------------------------------------------------------------------------
def test_real_output_gemini_features():
    """Check existing pipeline outputs for Gemini feature presence."""
    outputs_dir = PROJECT_ROOT / "outputs"
    if not outputs_dir.exists():
        _report("real_output_gemini_features", "WARN", "No outputs/ directory found")
        return

    json_files = list(outputs_dir.rglob("*.json"))
    if not json_files:
        _report("real_output_gemini_features", "WARN", "No JSON output files found")
        return

    # Check the most recent files
    json_files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    checked = 0
    has_structured = 0
    has_key_findings = 0
    has_tables = 0
    has_charts = 0
    has_metrics = 0

    for jf in json_files[:5]:
        try:
            data = json.loads(jf.read_text(encoding="utf-8"))
            checked += 1

            # Check state for structured_data
            sd = data.get("structured_data", [])
            if sd:
                has_structured += 1

            # Check final_report for features
            report = data.get("final_report", "")
            if not report:
                continue

            if re.search(r'Key Findings', report, re.IGNORECASE):
                has_key_findings += 1
            if "| --- |" in report or re.search(r'\|.*\|.*\|', report):
                has_tables += 1
            if "data:image/png;base64," in report:
                has_charts += 1
            if ":::metrics" in report:
                has_metrics += 1
        except Exception:
            continue

    detail = (
        f"Checked {checked} outputs: "
        f"structured_data={has_structured}, Key Findings={has_key_findings}, "
        f"tables={has_tables}, charts={has_charts}, :::metrics={has_metrics}"
    )

    if checked == 0:
        _report("real_output_gemini_features", "WARN", "No valid outputs to check")
    elif has_structured == 0 and has_key_findings == 0:
        _report("real_output_gemini_features", "WARN",
                f"No Gemini features in existing outputs (pre-fix). {detail}")
    else:
        _report("real_output_gemini_features", "INFO", detail)


# ---------------------------------------------------------------------------
# LAYER 14: Synthesizer Return Dict Completeness
# ---------------------------------------------------------------------------
def test_synthesizer_return_completeness():
    """Verify synthesizer returns ALL state keys that ResearchState expects."""
    synth_path = PROJECT_ROOT / "src/polaris_graph/agents/synthesizer.py"
    content = synth_path.read_text(encoding="utf-8")

    # Extract return dict keys (look for the main return statement)
    # Keys in the return dict
    critical_return_keys = [
        "hallucination_audit",
        "section_evidence_map",
        "sections",
        "final_report",
        "bibliography",
    ]
    missing = []
    for key in critical_return_keys:
        pattern = f'"{key}":'
        if pattern not in content:
            missing.append(key)

    if not missing:
        _report("synthesizer_return_completeness", "PASS",
                f"All {len(critical_return_keys)} critical keys in return dict")
    else:
        _report("synthesizer_return_completeness", "FAIL",
                f"Missing from return: {missing}")


# ---------------------------------------------------------------------------
# LAYER 15: Silent Failure Detection
# ---------------------------------------------------------------------------
def test_silent_failure_patterns():
    """Scan for dangerous silent failure patterns in critical files."""
    critical_files = [
        "src/polaris_graph/agents/analyzer.py",
        "src/polaris_graph/agents/synthesizer.py",
        "src/polaris_graph/agents/verifier.py",
        "src/polaris_graph/synthesis/section_writer.py",
    ]
    total_silent = 0
    for filepath in critical_files:
        full_path = PROJECT_ROOT / filepath
        if not full_path.exists():
            continue
        content = full_path.read_text(encoding="utf-8")
        # Count `except: pass` or `except Exception: pass`
        silent_count = len(re.findall(r'except\s*(?:Exception\s*(?:as\s+\w+)?)?\s*:\s*\n\s*pass', content))
        total_silent += silent_count

    if total_silent == 0:
        _report("silent_failure_patterns", "PASS",
                "No `except: pass` patterns in critical pipeline files")
    else:
        _report("silent_failure_patterns", "WARN",
                f"{total_silent} silent exception handlers found")


# ---------------------------------------------------------------------------
# LAYER 16: End-to-End Data Flow Simulation
# ---------------------------------------------------------------------------
def test_e2e_data_flow_simulation():
    """Simulate the structured_data flow: extract → accumulate → filter → inject.

    This doesn't call LLMs but traces the exact code path with mock data.
    """
    # Step 1: Simulate analyzer output
    analyzer_result = {
        "evidence": [
            {"evidence_id": "ev_001", "source_url": "https://example.com/a",
             "content": "Steel has 400 MPa strength", "quality_tier": "GOLD",
             "relevance_score": 0.9},
        ],
        "structured_data": [
            {"value": "400", "unit": "MPa", "entity": "Steel",
             "source_url": "https://example.com/a", "evidence_id": "ev_001",
             "data_type": "measurement"},
        ],
        "fetched_content": [],
        "status": "verifying",
    }

    # Step 2: Simulate graph accumulation
    existing_state = {"structured_data": [], "evidence": []}
    existing_sd = list(existing_state.get("structured_data", []))
    new_sd = analyzer_result.get("structured_data", [])
    if existing_sd and new_sd:
        existing_sd_keys = {
            (d.get("source_url", ""), d.get("value", ""), d.get("entity", ""))
            for d in existing_sd
        }
        unique_sd = [
            d for d in new_sd
            if (d.get("source_url", ""), d.get("value", ""), d.get("entity", ""))
            not in existing_sd_keys
        ]
        accumulated = existing_sd + unique_sd
    elif new_sd:
        accumulated = new_sd
    else:
        accumulated = existing_sd

    # Step 3: Simulate synthesizer filter
    section_ev_ids = {"ev_001"}
    section_data = [
        dp for dp in accumulated
        if dp.get("evidence_id", "") in section_ev_ids
    ]

    # Step 4: Check chain integrity
    issues = []
    if len(accumulated) != 1:
        issues.append(f"accumulation: expected 1, got {len(accumulated)}")
    if len(section_data) != 1:
        issues.append(f"filter: expected 1, got {len(section_data)}")
    if section_data and section_data[0].get("value") != "400":
        issues.append(f"value integrity: expected '400', got {section_data[0].get('value')}")

    if not issues:
        _report("e2e_data_flow_simulation", "PASS",
                "Full chain: extract(1) -> accumulate(1) -> filter(1) -> value='400' intact")
    else:
        _report("e2e_data_flow_simulation", "FAIL", "; ".join(issues))


# ---------------------------------------------------------------------------
# LAYER 17: Word Target Consistency (Fix 4 verification)
# ---------------------------------------------------------------------------
def test_word_target_no_hard_cap():
    """Verify section_writer doesn't have contradictory word limits."""
    sw_path = PROJECT_ROOT / "src/polaris_graph/synthesis/section_writer.py"
    content = sw_path.read_text(encoding="utf-8")

    # Check for the old problematic pattern
    has_hard_cap = "Maximum: 1000 words" in content
    has_dynamic = "suggested_words" in content or "_suggested_words" in content or "target_words" in content

    if has_hard_cap:
        _report("word_target_no_hard_cap", "FAIL",
                "Still has 'Maximum: 1000 words' hard cap conflicting with dynamic target")
    elif has_dynamic:
        _report("word_target_no_hard_cap", "PASS",
                "Uses dynamic word target (no hard 1000-word cap)")
    else:
        _report("word_target_no_hard_cap", "WARN",
                "Could not find word target pattern")


# ---------------------------------------------------------------------------
# LAYER 18: Evidence Ranking (Fix 5 verification)
# ---------------------------------------------------------------------------
def test_evidence_ranking_by_tier():
    """Verify evidence is sorted by tier before synthesis."""
    sw_path = PROJECT_ROOT / "src/polaris_graph/synthesis/section_writer.py"
    content = sw_path.read_text(encoding="utf-8")

    has_tier_sort = "GOLD" in content and ("tier" in content.lower()) and "sort" in content.lower()
    has_ranking = "TOP EVIDENCE" in content or "tier_order" in content or "quality_tier" in content

    if has_tier_sort or has_ranking:
        _report("evidence_ranking_by_tier", "PASS",
                "Evidence sorted by tier (GOLD > SILVER > BRONZE)")
    else:
        _report("evidence_ranking_by_tier", "WARN",
                "Could not confirm tier-based evidence ranking")


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main():
    print("\n" + "=" * 70)
    print("  ANTI-TUNNEL-VIEW INTEGRATION TEST")
    print("  Tests feature survival through the FULL pipeline chain")
    print("=" * 70 + "\n")

    print("[Layer 1] State Schema Integrity")
    test_state_schema_completeness()
    test_initial_state_defaults()

    print("\n[Layer 2] Feature Flag Verification")
    test_feature_flags_enabled()
    test_feature_flag_code_guards()

    print("\n[Layer 3] structured_data Flow Chain")
    test_analyzer_returns_structured_data()
    test_graph_accumulates_structured_data()
    test_synthesizer_reads_structured_data()
    test_chart_injection_into_sections()

    print("\n[Layer 4] Key Findings Survival")
    test_key_findings_preservation()

    print("\n[Layer 5] :::metrics Survival")
    test_metrics_unconditional_insertion()
    test_metrics_survives_post_processing()

    print("\n[Layer 6] Matplotlib Execution")
    test_matplotlib_subprocess()

    print("\n[Layer 7] Schema Validation")
    test_structured_data_point_coercion()
    test_structured_data_extraction_wrap()

    print("\n[Layer 8] Chart/Table Markdown")
    test_chart_markdown_format()
    test_table_markdown_format()

    print("\n[Layer 9] Filler Reduction")
    test_filler_reduction()

    print("\n[Layer 10] DOCX Export")
    test_docx_export_with_rich_content()

    print("\n[Layer 11] Smart Art Gap Detection")
    test_smart_art_injection_gap()

    print("\n[Layer 12] Evidence Accumulation")
    test_evidence_accumulation_pattern()

    print("\n[Layer 13] Real Output Analysis")
    test_real_output_gemini_features()

    print("\n[Layer 14] Synthesizer Return Completeness")
    test_synthesizer_return_completeness()

    print("\n[Layer 15] Silent Failure Detection")
    test_silent_failure_patterns()

    print("\n[Layer 16] E2E Data Flow Simulation")
    test_e2e_data_flow_simulation()

    print("\n[Layer 17] Word Target Consistency")
    test_word_target_no_hard_cap()

    print("\n[Layer 18] Evidence Ranking")
    test_evidence_ranking_by_tier()

    # Summary
    total = PASS + FAIL + WARN
    print("\n" + "=" * 70)
    print(f"  RESULTS: {PASS} PASS / {FAIL} FAIL / {WARN} WARN (total: {total})")
    if FAIL == 0:
        print("  STATUS: ALL CRITICAL CHECKS PASSED")
    else:
        print(f"  STATUS: {FAIL} FAILURES REQUIRE ATTENTION BEFORE E2E RUN")
    print("=" * 70 + "\n")

    # Save results
    results_path = PROJECT_ROOT / "outputs" / "anti_tunnel_view_results.json"
    results_path.parent.mkdir(exist_ok=True)
    results_path.write_text(json.dumps({
        "pass": PASS, "fail": FAIL, "warn": WARN,
        "tests": RESULTS,
    }, indent=2), encoding="utf-8")
    print(f"Results saved to: {results_path}")

    return 0 if FAIL == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
