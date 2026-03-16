"""
DOCX exporter for polaris graph research reports.

Converts a ResearchState result dict (as stored in outputs/polaris_graph/*.json)
into a professionally styled Microsoft Word (.docx) document with corporate
typography, inline superscript citations, bibliography, quality summary,
and audit certificate.

Requires: python-docx >= 1.1.0

LAW VI: All configurable parameters sourced from environment variables.
LAW V:  snake_case throughout, one-responsibility module.
"""

import base64
import hashlib
import io
import json
import logging
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Configuration from environment (LAW VI)
# ---------------------------------------------------------------------------

DOCX_BODY_FONT = os.getenv("POLARIS_DOCX_FONT", "Calibri")
DOCX_BODY_FONT_SIZE_PT = int(os.getenv("POLARIS_DOCX_FONT_SIZE", "11"))
DOCX_HEADING_FONT = os.getenv("POLARIS_DOCX_HEADING_FONT", "Calibri")
DOCX_HEADING1_SIZE_PT = int(os.getenv("POLARIS_DOCX_HEADING1_SIZE", "18"))
DOCX_HEADING2_SIZE_PT = int(os.getenv("POLARIS_DOCX_HEADING2_SIZE", "14"))
DOCX_HEADING3_SIZE_PT = int(os.getenv("POLARIS_DOCX_HEADING3_SIZE", "12"))
DOCX_TITLE_SIZE_PT = int(os.getenv("POLARIS_DOCX_TITLE_SIZE", "26"))
DOCX_MARGIN_CM = float(os.getenv("POLARIS_DOCX_MARGIN_CM", "2.54"))
DOCX_ACCENT_COLOR = os.getenv("POLARIS_DOCX_ACCENT_COLOR", "1F4E79")
DOCX_PIPELINE_VERSION = os.getenv("POLARIS_PIPELINE_VERSION", "polaris-graph-v1")
DOCX_LINE_SPACING = float(os.getenv("POLARIS_DOCX_LINE_SPACING", "1.15"))


# ---------------------------------------------------------------------------
# Markdown inline patterns
# ---------------------------------------------------------------------------

# Matches **bold** text
_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
# Matches *italic* text (but not **bold**)
_ITALIC_PATTERN = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
# Matches inline citations like [1], [23], [1][2]
_CITATION_PATTERN = re.compile(r"\[(\d+)\]")
# Matches markdown headings at start of line
_HEADING_PATTERN = re.compile(r"^(#{1,3})\s+(.+)$", re.MULTILINE)
# Matches markdown bullet list items
_BULLET_PATTERN = re.compile(r"^\s*[-*+]\s+(.+)$", re.MULTILINE)
# Matches markdown numbered list items
_NUMBERED_PATTERN = re.compile(r"^\s*\d+\.\s+(.+)$", re.MULTILINE)
# Matches markdown table separator row (e.g., | --- | --- |)
_TABLE_SEP_PATTERN = re.compile(r"^\s*\|[\s\-:]+(\|[\s\-:]+)+\|?\s*$")
# Matches markdown table row (e.g., | cell1 | cell2 |)
_TABLE_ROW_PATTERN = re.compile(r"^\s*\|(.+)\|\s*$")
# Matches base64 image in markdown (e.g., ![alt](data:image/png;base64,...))
_BASE64_IMG_PATTERN = re.compile(
    r"!\[([^\]]*)\]\(data:image/(\w+);base64,([A-Za-z0-9+/=\s]+)\)"
)
# Matches figure caption (e.g., *Figure 1: description*)
_FIGURE_CAPTION_PATTERN = re.compile(r"^\s*\*(.+)\*\s*$")


def _parse_accent_color(hex_str: str) -> RGBColor:
    """Parse a 6-char hex string into an RGBColor. Falls back to dark blue."""
    try:
        hex_str = hex_str.lstrip("#")
        if len(hex_str) != 6:
            raise ValueError(f"Invalid hex color length: {hex_str}")
        return RGBColor(
            int(hex_str[0:2], 16),
            int(hex_str[2:4], 16),
            int(hex_str[4:6], 16),
        )
    except (ValueError, IndexError):
        logger.warning(
            "Invalid POLARIS_DOCX_ACCENT_COLOR '%s', falling back to #1F4E79",
            hex_str,
        )
        return RGBColor(0x1F, 0x4E, 0x79)


class DocxExporter:
    """
    Generates professionally styled Microsoft Word documents from
    POLARIS research pipeline output.

    Usage::

        exporter = DocxExporter()
        output_path = exporter.export(report_data, Path("outputs/report.docx"))
    """

    def __init__(self) -> None:
        self._accent_color = _parse_accent_color(DOCX_ACCENT_COLOR)
        self._body_font = DOCX_BODY_FONT
        self._body_size = Pt(DOCX_BODY_FONT_SIZE_PT)
        self._heading_font = DOCX_HEADING_FONT
        self._line_spacing = DOCX_LINE_SPACING
        self._margin = Cm(DOCX_MARGIN_CM)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def export(self, report_data: dict, output_path: Path) -> Path:
        """
        Export a POLARIS research result dict to a styled .docx file.

        Args:
            report_data: The full ResearchState dict (deserialized JSON).
                         Expected keys: vector_id, original_query (or query),
                         status, final_report, bibliography, quality_metrics,
                         sections, evidence (or evidence_count),
                         iteration_count, timestamps.
            output_path:  Destination path for the .docx file. Parent
                         directory is created if it does not exist.

        Returns:
            The resolved Path where the document was written.

        Raises:
            ValueError: If report_data is missing critical keys.
            OSError:    If the file cannot be written.
        """
        output_path = Path(output_path)
        self._validate_report_data(report_data)

        doc = Document()
        self._configure_styles(doc)
        self._set_margins(doc)

        # --- Title Page ---
        self._add_title_page(doc, report_data)
        doc.add_page_break()

        # --- Table of Contents ---
        self._add_table_of_contents(doc)
        doc.add_page_break()

        # --- Report Body ---
        self._add_report_body(doc, report_data)
        doc.add_page_break()

        # --- Bibliography ---
        self._add_bibliography(doc, report_data)
        doc.add_page_break()

        # --- Quality Summary ---
        self._add_quality_summary(doc, report_data)
        doc.add_page_break()

        # --- Audit Certificate ---
        self._add_audit_certificate(doc, report_data)

        # Write to disk
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(
            "DOCX exported: %s (%d bytes)",
            output_path,
            output_path.stat().st_size,
        )
        return output_path

    # ------------------------------------------------------------------
    # Validation
    # ------------------------------------------------------------------

    @staticmethod
    def _validate_report_data(report_data: dict) -> None:
        """Ensure minimum required keys are present."""
        required_keys = {"vector_id", "status"}
        missing = required_keys - set(report_data.keys())
        if missing:
            raise ValueError(
                f"report_data missing required keys: {missing}"
            )
        # At least one of these must be present for the report body
        has_body = bool(
            report_data.get("final_report")
            or report_data.get("sections")
        )
        if not has_body:
            raise ValueError(
                "report_data must contain 'final_report' or 'sections'"
            )

    # ------------------------------------------------------------------
    # Style configuration
    # ------------------------------------------------------------------

    def _configure_styles(self, doc: Document) -> None:
        """Set up corporate typography for all built-in styles."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = self._body_font
        font.size = self._body_size
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = self._line_spacing

        # Heading styles
        heading_configs = [
            ("Heading 1", DOCX_HEADING1_SIZE_PT, True),
            ("Heading 2", DOCX_HEADING2_SIZE_PT, True),
            ("Heading 3", DOCX_HEADING3_SIZE_PT, False),
        ]
        for style_name, size_pt, bold in heading_configs:
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
            else:
                h_style = doc.styles.add_style(
                    style_name, WD_STYLE_TYPE.PARAGRAPH
                )
            h_font = h_style.font
            h_font.name = self._heading_font
            h_font.size = Pt(size_pt)
            h_font.bold = bold
            h_font.color.rgb = self._accent_color
            h_pf = h_style.paragraph_format
            h_pf.space_before = Pt(18)
            h_pf.space_after = Pt(8)
            h_pf.keep_with_next = True

        # Title style
        if "Title" in doc.styles:
            t_style = doc.styles["Title"]
            t_font = t_style.font
            t_font.name = self._heading_font
            t_font.size = Pt(DOCX_TITLE_SIZE_PT)
            t_font.bold = True
            t_font.color.rgb = self._accent_color

    def _set_margins(self, doc: Document) -> None:
        """Apply uniform margins to all sections."""
        for section in doc.sections:
            section.top_margin = self._margin
            section.bottom_margin = self._margin
            section.left_margin = self._margin
            section.right_margin = self._margin

    # ------------------------------------------------------------------
    # Title Page
    # ------------------------------------------------------------------

    def _add_title_page(self, doc: Document, report_data: dict) -> None:
        """Build a title page with query, timestamp, and key metrics."""
        # Spacer
        for _ in range(4):
            doc.add_paragraph("")

        # Title
        query = report_data.get("original_query") or report_data.get(
            "query", "Untitled Research"
        )
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(query)
        title_run.font.name = self._heading_font
        title_run.font.size = Pt(DOCX_TITLE_SIZE_PT)
        title_run.font.bold = True
        title_run.font.color.rgb = self._accent_color

        # Subtitle line
        doc.add_paragraph("")
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run("POLARIS Research Report")
        sub_run.font.name = self._heading_font
        sub_run.font.size = Pt(16)
        sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Horizontal rule (thin line)
        doc.add_paragraph("")
        hr_para = doc.add_paragraph()
        hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr_run = hr_para.add_run("_" * 60)
        hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        hr_run.font.size = Pt(8)

        doc.add_paragraph("")

        # Metadata block
        timestamps = report_data.get("timestamps", {})
        created_raw = timestamps.get("created", "")
        generation_ts = self._format_timestamp(created_raw)

        quality_metrics = report_data.get("quality_metrics") or {}
        faithfulness = quality_metrics.get(
            "faithfulness_score",
            report_data.get("faithfulness_score", 0.0),
        )
        evidence_count = self._resolve_evidence_count(report_data)
        source_count = quality_metrics.get(
            "unique_sources",
            self._count_unique_sources(report_data),
        )

        meta_lines = [
            ("Generated", generation_ts),
            ("Vector ID", report_data.get("vector_id", "N/A")),
            ("Faithfulness", f"{faithfulness * 100:.1f}%"),
            ("Evidence Pieces", str(evidence_count)),
            ("Unique Sources", str(source_count)),
        ]

        for label, value in meta_lines:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = meta_para.add_run(f"{label}: ")
            label_run.font.name = self._body_font
            label_run.font.size = Pt(12)
            label_run.font.bold = True
            label_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            value_run = meta_para.add_run(value)
            value_run.font.name = self._body_font
            value_run.font.size = Pt(12)
            value_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ------------------------------------------------------------------
    # Table of Contents (Word field code)
    # ------------------------------------------------------------------

    @staticmethod
    def _add_table_of_contents(doc: Document) -> None:
        """
        Insert a Table of Contents placeholder using Word field codes.

        The TOC is not populated by python-docx itself; the user must
        right-click the field in Word and select 'Update Field' (or
        press Ctrl+A then F9) to generate page numbers.
        """
        toc_heading = doc.add_heading("Table of Contents", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = r' TOC \o "1-3" \h \z \u '
        run._r.append(instr_text)

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_char_separate)

        # Placeholder text visible before user updates the field
        placeholder_run = paragraph.add_run(
            "Right-click here and select 'Update Field' to generate "
            "Table of Contents."
        )
        placeholder_run.font.italic = True
        placeholder_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        fld_char_end_run = paragraph.add_run()
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        fld_char_end_run._r.append(fld_char_end)

    # ------------------------------------------------------------------
    # Report Body
    # ------------------------------------------------------------------

    def _add_report_body(self, doc: Document, report_data: dict) -> None:
        """
        Render the report body into the document.

        Strategy:
        1. If structured ``sections`` are available, render each with
           its heading and parsed content.
        2. Otherwise, parse the ``final_report`` markdown string.
        """
        sections = report_data.get("sections")
        if sections and isinstance(sections, list) and len(sections) > 0:
            self._render_structured_sections(doc, sections)
        else:
            final_report = report_data.get("final_report", "")
            if final_report:
                self._render_markdown(doc, final_report)

    def _render_structured_sections(
        self, doc: Document, sections: list[dict]
    ) -> None:
        """Render pre-structured section dicts with heading + content."""
        for section in sections:
            title = section.get("title", "Untitled Section")
            content = section.get("content", "")
            doc.add_heading(title, level=2)
            self._render_markdown(doc, content)

    def _render_markdown(self, doc: Document, text: str) -> None:
        """
        Parse a markdown string and append paragraphs to the document.

        Handles: headings (# ## ###), bold (**), italic (*),
        bullet lists (- * +), numbered lists (1.), inline citations [N].
        """
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Heading
            heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                doc.add_heading(heading_text, level=level)
                i += 1
                continue

            # Bullet list item
            bullet_match = re.match(r"^\s*[-*+]\s+(.+)$", stripped)
            if bullet_match:
                item_text = bullet_match.group(1)
                para = doc.add_paragraph(style="List Bullet")
                self._add_rich_text(para, item_text)
                i += 1
                continue

            # Numbered list item
            num_match = re.match(r"^\s*\d+\.\s+(.+)$", stripped)
            if num_match:
                item_text = num_match.group(1)
                para = doc.add_paragraph(style="List Number")
                self._add_rich_text(para, item_text)
                i += 1
                continue

            # Markdown table detection: look ahead for separator row
            if _TABLE_ROW_PATTERN.match(stripped):
                # Check if next line is a separator row (header | --- | --- |)
                if i + 1 < len(lines) and _TABLE_SEP_PATTERN.match(
                    lines[i + 1].strip()
                ):
                    # Collect all table rows
                    table_lines = [stripped]
                    i += 1
                    while i < len(lines):
                        row_stripped = lines[i].strip()
                        if not row_stripped:
                            break
                        if not _TABLE_ROW_PATTERN.match(
                            row_stripped
                        ) and not _TABLE_SEP_PATTERN.match(row_stripped):
                            break
                        table_lines.append(row_stripped)
                        i += 1
                    self._render_markdown_table(doc, table_lines)
                    continue

            # Base64 image detection
            img_match = _BASE64_IMG_PATTERN.search(stripped)
            if img_match:
                self._render_base64_image(
                    doc,
                    alt_text=img_match.group(1),
                    img_format=img_match.group(2),
                    b64_data=img_match.group(3),
                )
                i += 1
                continue

            # Key Findings block detection (MUST be before figure caption
            # because **Key Findings:** matches the *...* caption pattern)
            # FIX-7: Also detect ## Key Findings and ### Key Findings headings
            _is_key_findings = (
                stripped in ("**Key Findings:**", "**Key Findings**")
                or stripped.startswith("**Key Findings")
                or bool(re.match(r'^#{2,4}\s+Key Findings', stripped, re.IGNORECASE))
            )
            if _is_key_findings:
                # Render as styled heading-like block, then collect bullets
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.left_indent = Cm(0.5)
                run = para.add_run("Key Findings")
                run.font.name = self._heading_font
                run.font.size = Pt(DOCX_HEADING3_SIZE_PT)
                run.font.bold = True
                run.font.color.rgb = self._accent_color
                # Add left border via XML (accent color bar)
                self._add_left_border(para, self._accent_color)
                i += 1
                # Collect following bullet items as Key Findings
                while i < len(lines):
                    kf_line = lines[i].strip()
                    if not kf_line:
                        i += 1
                        continue
                    kf_bullet = re.match(r"^\s*[-*+]\s+(.+)$", kf_line)
                    if kf_bullet:
                        kf_para = doc.add_paragraph(style="List Bullet")
                        kf_para.paragraph_format.left_indent = Cm(1.0)
                        self._add_rich_text(kf_para, kf_bullet.group(1))
                        i += 1
                    else:
                        break
                continue

            # Figure caption (italic line like *Figure 1: ...*  or *Table 1: ...*)
            caption_match = _FIGURE_CAPTION_PATTERN.match(stripped)
            if caption_match:
                caption_text = caption_match.group(1)
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(12)
                run = para.add_run(caption_text)
                run.font.name = self._body_font
                run.font.size = Pt(max(DOCX_BODY_FONT_SIZE_PT - 2, 8))
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                i += 1
                continue

            # Regular paragraph -- accumulate continuation lines
            paragraph_lines = [stripped]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    break
                # Stop if next line is a heading, bullet, numbered item,
                # table row, or image
                if re.match(r"^#{1,3}\s+", next_line):
                    break
                if re.match(r"^\s*[-*+]\s+", next_line):
                    break
                if re.match(r"^\s*\d+\.\s+", next_line):
                    break
                if _TABLE_ROW_PATTERN.match(next_line):
                    break
                if _BASE64_IMG_PATTERN.search(next_line):
                    break
                paragraph_lines.append(next_line)
                i += 1

            full_text = " ".join(paragraph_lines)
            para = doc.add_paragraph()
            self._add_rich_text(para, full_text)

    def _add_rich_text(self, paragraph, text: str) -> None:
        """
        Parse inline markdown (bold, italic, citations) and append
        styled runs to the given paragraph.

        Citations like [1] are rendered as superscript.
        **bold** is rendered bold. *italic* is rendered italic.
        """
        # Tokenize the text into segments: plain, bold, italic, citation
        # We use a combined regex that matches any of the three patterns.
        combined_pattern = re.compile(
            r"(\*\*(.+?)\*\*)"      # group 1=full bold, group 2=bold content
            r"|"
            r"((?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*))"  # group 3=full italic, group 4=italic content
            r"|"
            r"(\[(\d+)\])"          # group 5=full citation, group 6=number
        )

        last_end = 0
        for match in combined_pattern.finditer(text):
            # Add any plain text before this match
            start = match.start()
            if start > last_end:
                plain = text[last_end:start]
                if plain:
                    run = paragraph.add_run(plain)
                    run.font.name = self._body_font
                    run.font.size = self._body_size

            if match.group(2) is not None:
                # Bold
                run = paragraph.add_run(match.group(2))
                run.font.name = self._body_font
                run.font.size = self._body_size
                run.font.bold = True
            elif match.group(4) is not None:
                # Italic
                run = paragraph.add_run(match.group(4))
                run.font.name = self._body_font
                run.font.size = self._body_size
                run.font.italic = True
            elif match.group(6) is not None:
                # Citation as superscript
                run = paragraph.add_run(f"[{match.group(6)}]")
                run.font.name = self._body_font
                run.font.size = Pt(max(DOCX_BODY_FONT_SIZE_PT - 3, 7))
                run.font.superscript = True
                run.font.color.rgb = self._accent_color

            last_end = match.end()

        # Trailing plain text
        if last_end < len(text):
            remaining = text[last_end:]
            if remaining:
                run = paragraph.add_run(remaining)
                run.font.name = self._body_font
                run.font.size = self._body_size

    # ------------------------------------------------------------------
    # Markdown Table Rendering
    # ------------------------------------------------------------------

    def _render_markdown_table(
        self, doc: Document, table_lines: list[str]
    ) -> None:
        """
        Parse pipe-delimited markdown table lines into a python-docx Table
        with alternating row shading.

        Expected format:
            | Header1 | Header2 |
            | ------- | ------- |
            | Cell1   | Cell2   |
        """
        # Parse rows, skipping separator lines
        parsed_rows: list[list[str]] = []
        for line in table_lines:
            if _TABLE_SEP_PATTERN.match(line):
                continue
            cells = [
                c.strip()
                for c in line.strip().strip("|").split("|")
            ]
            parsed_rows.append(cells)

        if len(parsed_rows) < 2:
            # Need at least header + 1 data row
            return

        headers = parsed_rows[0]
        data_rows = parsed_rows[1:]
        col_count = len(headers)

        table = doc.add_table(
            rows=1 + len(data_rows), cols=col_count
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for c_idx, header_text in enumerate(headers):
            cell = table.cell(0, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(header_text)
            run.font.name = self._heading_font
            run.font.size = self._body_size
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            self._shade_cell(cell, self._accent_color)

        # Data rows with alternating shading
        stripe_color = RGBColor(0xF2, 0xF2, 0xF2)
        for r_idx, row_cells in enumerate(data_rows):
            for c_idx in range(col_count):
                cell_text = (
                    row_cells[c_idx] if c_idx < len(row_cells) else ""
                )
                cell = table.cell(r_idx + 1, c_idx)
                cell.text = ""
                para = cell.paragraphs[0]
                self._add_rich_text(para, cell_text)
                # Alternating row shading (even data rows)
                if r_idx % 2 == 1:
                    self._shade_cell(cell, stripe_color)

        # Add spacing after table
        doc.add_paragraph("")

    # ------------------------------------------------------------------
    # Base64 Image Rendering
    # ------------------------------------------------------------------

    def _render_base64_image(
        self,
        doc: Document,
        alt_text: str,
        img_format: str,
        b64_data: str,
    ) -> None:
        """
        Decode a base64-encoded image and embed it inline in the document.

        Args:
            doc: The python-docx Document.
            alt_text: Alt text from the markdown image tag.
            img_format: Image format (png, jpg, etc.).
            b64_data: Base64-encoded image data.
        """
        try:
            # Strip whitespace from base64 data
            clean_b64 = b64_data.replace("\n", "").replace(" ", "")
            image_bytes = base64.b64decode(clean_b64)
            image_stream = io.BytesIO(image_bytes)

            # Add image paragraph (centered)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)

            run = para.add_run()
            run.add_picture(image_stream, width=Inches(5.5))

            logger.debug(
                "DOCX: embedded base64 image '%s' (%d bytes)",
                alt_text,
                len(image_bytes),
            )

        except Exception as exc:
            logger.warning(
                "DOCX: failed to embed base64 image '%s': %s",
                alt_text,
                str(exc)[:200],
            )
            # Fallback: add placeholder text
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"[Image: {alt_text}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ------------------------------------------------------------------
    # Paragraph Border Helper
    # ------------------------------------------------------------------

    @staticmethod
    def _add_left_border(
        paragraph, color: RGBColor, width_pt: int = 4
    ) -> None:
        """Add a left border (accent bar) to a paragraph via XML."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(width_pt * 2))  # half-points
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), f"{color}")
        pBdr.append(left)
        pPr.append(pBdr)

    # ------------------------------------------------------------------
    # Bibliography
    # ------------------------------------------------------------------

    def _add_bibliography(self, doc: Document, report_data: dict) -> None:
        """Render the bibliography as a numbered list with title + URL."""
        doc.add_heading("Bibliography", level=1)

        bibliography = report_data.get("bibliography", [])
        if not bibliography:
            para = doc.add_paragraph("No bibliography entries available.")
            para.font = self._body_font
            return

        for idx, entry in enumerate(bibliography, start=1):
            citation_key = entry.get("citation_key", f"[{idx}]")
            # Normalize citation key to just the number
            num_match = re.search(r"\d+", str(citation_key))
            display_num = num_match.group(0) if num_match else str(idx)

            title = entry.get("title", "")
            url = entry.get("url", "")
            formatted = entry.get("formatted", "")
            quality_tier = entry.get("quality_tier", entry.get("source_type", ""))

            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.left_indent = Cm(1.0)
            para.paragraph_format.first_line_indent = Cm(-1.0)

            # Number prefix
            num_run = para.add_run(f"[{display_num}] ")
            num_run.font.name = self._body_font
            num_run.font.size = self._body_size
            num_run.font.bold = True
            num_run.font.color.rgb = self._accent_color

            # Use formatted citation if available, otherwise construct
            if formatted:
                text_run = para.add_run(formatted)
                text_run.font.name = self._body_font
                text_run.font.size = self._body_size
            elif title:
                title_run = para.add_run(title)
                title_run.font.name = self._body_font
                title_run.font.size = self._body_size
                title_run.font.italic = True

            # URL on separate line if present
            if url:
                if formatted or title:
                    para.add_run(" ")
                url_run = para.add_run(url)
                url_run.font.name = self._body_font
                url_run.font.size = Pt(max(DOCX_BODY_FONT_SIZE_PT - 1, 8))
                url_run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

    # ------------------------------------------------------------------
    # Quality Summary
    # ------------------------------------------------------------------

    def _add_quality_summary(
        self, doc: Document, report_data: dict
    ) -> None:
        """Render a quality summary page with a metrics table."""
        doc.add_heading("Quality Summary", level=1)

        quality_metrics = report_data.get("quality_metrics") or {}
        faithfulness = quality_metrics.get(
            "faithfulness_score",
            report_data.get("faithfulness_score", 0.0),
        )
        evidence_count = self._resolve_evidence_count(report_data)
        source_count = quality_metrics.get(
            "unique_sources",
            self._count_unique_sources(report_data),
        )
        iterations = report_data.get("iteration_count", 0)
        word_count = quality_metrics.get("total_words", 0)
        total_sections = quality_metrics.get("total_sections", 0)
        total_citations = quality_metrics.get("total_citations", 0)
        coverage = quality_metrics.get("coverage_score", 0.0)
        coherence = quality_metrics.get("coherence_score", 0.0)

        # Build metrics table
        rows = [
            ("Metric", "Value"),
            ("Faithfulness Score", f"{faithfulness * 100:.1f}%"),
            ("Evidence Count", str(evidence_count)),
            ("Unique Sources", str(source_count)),
            ("Pipeline Iterations", str(iterations)),
            ("Total Word Count", f"{word_count:,}"),
            ("Total Sections", str(total_sections)),
            ("Total Citations", str(total_citations)),
            ("Coverage Score", f"{coverage * 100:.1f}%"),
            ("Coherence Score", f"{coherence * 100:.1f}%"),
            ("Pipeline Status", report_data.get("status", "unknown")),
            (
                "Quality Gate",
                report_data.get("quality_gate_result", "N/A"),
            ),
        ]

        table = doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 1"

        for row_idx, (label, value) in enumerate(rows):
            cell_label = table.cell(row_idx, 0)
            cell_value = table.cell(row_idx, 1)
            cell_label.text = label
            cell_value.text = value

            # Style header row
            if row_idx == 0:
                for cell in (cell_label, cell_value):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                            run.font.name = self._heading_font
                            run.font.size = self._body_size
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    self._shade_cell(cell, self._accent_color)
            else:
                for cell in (cell_label, cell_value):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = self._body_font
                            run.font.size = self._body_size

        # Evidence tier breakdown if available
        gold = quality_metrics.get("gold_evidence", 0)
        silver = quality_metrics.get("silver_evidence", 0)
        bronze = evidence_count - gold - silver
        if evidence_count > 0:
            doc.add_paragraph("")
            tier_para = doc.add_heading("Evidence Tier Distribution", level=2)
            tier_rows = [
                ("Tier", "Count", "Percentage"),
                ("GOLD", str(gold), f"{gold / evidence_count * 100:.1f}%"),
                (
                    "SILVER",
                    str(silver),
                    f"{silver / evidence_count * 100:.1f}%",
                ),
                (
                    "BRONZE",
                    str(max(bronze, 0)),
                    f"{max(bronze, 0) / evidence_count * 100:.1f}%",
                ),
            ]
            tier_table = doc.add_table(rows=len(tier_rows), cols=3)
            tier_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            tier_table.style = "Light Grid Accent 1"
            for r_idx, row_data in enumerate(tier_rows):
                for c_idx, cell_text in enumerate(row_data):
                    cell = tier_table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    if r_idx == 0:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.name = self._heading_font
                                run.font.size = self._body_size
                                run.font.color.rgb = RGBColor(
                                    0xFF, 0xFF, 0xFF
                                )
                        self._shade_cell(cell, self._accent_color)
                    else:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = self._body_font
                                run.font.size = self._body_size

    # ------------------------------------------------------------------
    # Audit Certificate
    # ------------------------------------------------------------------

    def _add_audit_certificate(
        self, doc: Document, report_data: dict
    ) -> None:
        """
        Render an audit certificate with vector ID, SHA-256 hash of the
        report content, pipeline version, and generation timestamp.
        """
        doc.add_heading("Audit Certificate", level=1)

        vector_id = report_data.get("vector_id", "N/A")
        final_report = report_data.get("final_report", "")
        content_hash = hashlib.sha256(
            final_report.encode("utf-8")
        ).hexdigest()

        timestamps = report_data.get("timestamps", {})
        created = self._format_timestamp(timestamps.get("created", ""))
        completed = self._format_timestamp(
            timestamps.get("synthesize_end", timestamps.get("completed", ""))
        )

        certificate_fields = [
            ("Vector ID", vector_id),
            ("Content SHA-256", content_hash),
            ("Pipeline Version", DOCX_PIPELINE_VERSION),
            ("Generation Timestamp", created),
            ("Completion Timestamp", completed),
            ("Pipeline Status", report_data.get("status", "unknown")),
            (
                "Iterations Completed",
                str(report_data.get("iteration_count", 0)),
            ),
            (
                "Convergence Reason",
                report_data.get("convergence_reason", "N/A") or "N/A",
            ),
        ]

        # Render as a clean two-column table
        table = doc.add_table(rows=len(certificate_fields), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, (label, value) in enumerate(certificate_fields):
            cell_label = table.cell(row_idx, 0)
            cell_value = table.cell(row_idx, 1)

            label_para = cell_label.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.name = self._body_font
            label_run.font.size = self._body_size
            label_run.font.bold = True

            value_para = cell_value.paragraphs[0]
            # SHA-256 in monospace
            if "SHA-256" in label:
                value_run = value_para.add_run(value)
                value_run.font.name = "Consolas"
                value_run.font.size = Pt(9)
                value_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                value_run = value_para.add_run(value)
                value_run.font.name = self._body_font
                value_run.font.size = self._body_size

        # Add certification statement
        doc.add_paragraph("")
        cert_para = doc.add_paragraph()
        cert_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cert_run = cert_para.add_run(
            "This document was generated by the POLARIS automated research "
            "pipeline. The SHA-256 hash above can be used to verify the "
            "integrity of the report content."
        )
        cert_run.font.name = self._body_font
        cert_run.font.size = Pt(9)
        cert_run.font.italic = True
        cert_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _shade_cell(cell, color: RGBColor) -> None:
        """Apply a background shading color to a table cell."""
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), f"{color}")
        shading_elm.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def _format_timestamp(raw: str) -> str:
        """
        Format an ISO timestamp string into a human-readable form.
        Returns 'N/A' if the string is empty or unparseable.
        """
        if not raw:
            return "N/A"
        try:
            dt = datetime.fromisoformat(raw)
            return dt.strftime("%Y-%m-%d %H:%M:%S UTC")
        except (ValueError, TypeError):
            return raw

    @staticmethod
    def _resolve_evidence_count(report_data: dict) -> int:
        """
        Determine the evidence count from the best available source.

        Priority: quality_metrics.total_evidence > evidence_count key >
                  len(evidence list).
        """
        qm = report_data.get("quality_metrics") or {}
        if qm.get("total_evidence"):
            return int(qm["total_evidence"])
        if report_data.get("evidence_count"):
            return int(report_data["evidence_count"])
        evidence = report_data.get("evidence", [])
        if isinstance(evidence, list):
            return len(evidence)
        return 0

    @staticmethod
    def _count_unique_sources(report_data: dict) -> int:
        """
        Count unique sources from bibliography or evidence list.
        """
        bibliography = report_data.get("bibliography", [])
        if bibliography:
            return len(bibliography)
        evidence = report_data.get("evidence", [])
        if isinstance(evidence, list):
            urls = {
                e.get("source_url", "")
                for e in evidence
                if isinstance(e, dict) and e.get("source_url")
            }
            return len(urls)
        return 0
