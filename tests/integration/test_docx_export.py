"""
Integration tests for the DOCX exporter (polaris graph research reports).

Tests REAL python-docx document generation and parsing. Every test creates
a genuine .docx file on disk via ``tmp_path``, then reads it back with
python-docx to verify structure, styles, and content. Zero mocks.

When a real pipeline result exists at ``outputs/polaris_graph/``, the
``report_data`` fixture loads it so that tests exercise production-scale
data. Otherwise, a comprehensive fixture dict matching the full pipeline
schema is used.

Module under test:
    src/polaris_graph/export/docx_exporter.py  -- DocxExporter.export()

LAW II:  All assertions verify real file artifacts.
LAW VI: No hard-coded paths; ``tmp_path`` and ``POLARIS_OUTPUT_DIR`` used.
"""

import glob
import hashlib
import json
import os
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from src.polaris_graph.export.docx_exporter import DocxExporter


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

_PIPELINE_OUTPUT_DIR = Path(os.getenv(
    "POLARIS_OUTPUT_DIR",
    "C:/POLARIS/outputs/polaris_graph",
))


def _find_real_result() -> dict | None:
    """Attempt to load a real pipeline JSON result with all required fields."""
    pattern = str(_PIPELINE_OUTPUT_DIR / "PG_TEST_*.json")
    for filepath in sorted(glob.glob(pattern), reverse=True):
        try:
            with open(filepath, encoding="utf-8") as fh:
                data = json.load(fh)
            has_all = all([
                data.get("final_report"),
                data.get("sections"),
                data.get("bibliography"),
                data.get("quality_metrics"),
                data.get("vector_id"),
                data.get("status"),
            ])
            if has_all:
                return data
        except (json.JSONDecodeError, OSError):
            continue
    return None


@pytest.fixture
def exporter():
    """Fresh DocxExporter instance per test."""
    return DocxExporter()


@pytest.fixture
def report_data():
    """Comprehensive fixture matching the real pipeline output schema.

    Prefers a real pipeline result when available; falls back to a
    hand-crafted dict that exercises every code path in the exporter.
    """
    real = _find_real_result()
    if real is not None:
        return real

    return {
        "vector_id": "V_TEST_001",
        "original_query": (
            "What are the most effective PFAS water filtration "
            "technologies for municipal treatment plants?"
        ),
        "status": "completed",
        "final_report": (
            "# Introduction\n\n"
            "PFAS (per- and polyfluoroalkyl substances) represent a class of "
            "persistent environmental contaminants [1] that have been detected "
            "in drinking water supplies worldwide.\n\n"
            "## Key Filtration Methods\n\n"
            "- **Granular Activated Carbon (GAC)** adsorption remains the most "
            "widely deployed treatment [2]\n"
            "- Ion exchange resins provide selective PFAS removal [3]\n"
            "- Reverse osmosis membranes achieve >99% rejection rates [4]\n\n"
            "## Emerging Technologies\n\n"
            "Recent advances in *electrochemical oxidation* and biochar-based "
            "adsorbents offer promising alternatives for shorter-chain PFAS "
            "compounds that resist conventional treatment.\n\n"
            "## Cost-Effectiveness\n\n"
            "A 2024 lifecycle cost analysis demonstrated that GAC systems "
            "operating at 10-minute EBCT achieve the lowest cost per 1000 "
            "gallons treated for PFOA and PFOS removal [1][2].\n\n"
            "## Regulatory Landscape\n\n"
            "The EPA finalized individual maximum contaminant levels (MCLs) "
            "of 4.0 ppt for PFOA and PFOS in April 2024 [3], driving "
            "widespread municipal adoption of advanced treatment.\n\n"
            "## Conclusions\n\n"
            "Multiple treatment approaches show efficacy for PFAS removal, "
            "with the optimal technology depending on influent characteristics, "
            "target compounds, and operational budget.\n\n"
            "[1] Patterson et al. 2019. Effectiveness of POU/POE systems.\n"
            "[2] Herkert et al. 2020. Assessing PFAS removal by POU filters.\n"
            "[3] EPA 2024. PFAS National Primary Drinking Water Regulation.\n"
            "[4] MacKeown et al. 2024. Removal of PFAS by POU/POE devices.\n"
        ),
        "sections": [
            {
                "title": "Introduction",
                "content": (
                    "PFAS are persistent pollutants [1] with **significant** "
                    "health risks including immunotoxicity and cancer."
                ),
            },
            {
                "title": "Key Filtration Methods",
                "content": (
                    "GAC adsorption [2], ion exchange [3], and reverse "
                    "osmosis [4] are primary municipal treatment methods."
                ),
            },
            {
                "title": "Emerging Technologies",
                "content": (
                    "Electrochemical oxidation and biochar-based adsorbents "
                    "offer *promising* alternatives for short-chain PFAS."
                ),
            },
            {
                "title": "Cost-Effectiveness",
                "content": (
                    "Lifecycle cost analysis shows GAC at 10-min EBCT "
                    "achieves lowest cost per 1000 gallons [1][2]."
                ),
            },
            {
                "title": "Regulatory Landscape",
                "content": (
                    "EPA finalized MCLs of 4.0 ppt for PFOA/PFOS in "
                    "April 2024 [3], driving municipal adoption."
                ),
            },
            {
                "title": "Conclusions",
                "content": (
                    "Optimal technology depends on influent, target "
                    "compounds, and operational budget."
                ),
            },
        ],
        "bibliography": [
            {
                "citation_key": "[1]",
                "title": "Effectiveness of POU/POE systems for PFAS removal",
                "url": "https://pmc.ncbi.nlm.nih.gov/articles/PMC6650157/",
                "quality_tier": "GOLD",
            },
            {
                "citation_key": "[2]",
                "title": "Assessing PFAS removal by point-of-use filters",
                "url": "https://pubs.acs.org/doi/10.1021/acs.estlett.0c00004",
                "quality_tier": "GOLD",
            },
            {
                "citation_key": "[3]",
                "title": "PFAS National Primary Drinking Water Regulation",
                "url": "https://www.epa.gov/pfas/pfas-drinking-water-regulation",
                "quality_tier": "GOLD",
            },
            {
                "citation_key": "[4]",
                "title": "Removal of PFAS by POU/POE devices: a review",
                "url": "https://www.sciencedirect.com/science/article/pii/S0048969724069213",
                "quality_tier": "SILVER",
            },
        ],
        "quality_metrics": {
            "faithfulness_score": 0.85,
            "total_evidence": 150,
            "unique_sources": 25,
            "total_words": 8000,
            "total_sections": 6,
            "total_citations": 45,
            "coverage_score": 0.9,
            "coherence_score": 0.88,
            "gold_evidence": 20,
            "silver_evidence": 50,
        },
        "timestamps": {"created": "2024-01-15T10:30:00"},
        "iteration_count": 3,
        "convergence_reason": "quality_gate_passed",
        "quality_gate_result": "PASS",
        "faithfulness_score": 0.85,
        "evidence_count": 150,
    }


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _export_and_load(exporter: DocxExporter, data: dict, tmp_path: Path):
    """Export to .docx then load back as a python-docx Document."""
    output_path = tmp_path / f"{data.get('vector_id', 'test')}.docx"
    result_path = exporter.export(data, output_path)
    doc = DocxDocument(str(result_path))
    return result_path, doc


def _full_text(doc: DocxDocument) -> str:
    """Extract all paragraph text from a Document, newline-joined."""
    return "\n".join(p.text for p in doc.paragraphs)


def _all_text(doc: DocxDocument) -> str:
    """Extract ALL text from a Document including paragraph and table cell content.

    python-docx separates paragraphs and tables at the document body level.
    Metrics tables (Quality Summary, Audit Certificate) render values inside
    table cells which are not included in ``doc.paragraphs``.  This helper
    collects text from both sources.
    """
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 1-2: Basic file generation and validity
# ---------------------------------------------------------------------------


class TestBasicExport:
    """Verify .docx file creation and structural validity."""

    def test_export_creates_valid_docx(self, exporter, report_data, tmp_path):
        """Export must create a non-empty .docx file on disk."""
        output_path = tmp_path / "test_report.docx"
        result_path = exporter.export(report_data, output_path)

        assert result_path.exists(), f"File not found: {result_path}"
        assert result_path.stat().st_size > 0, "File must be non-empty"
        assert result_path.suffix == ".docx"

    def test_docx_loads_with_python_docx(self, exporter, report_data, tmp_path):
        """The generated file must be a valid OOXML document loadable by python-docx."""
        result_path, doc = _export_and_load(exporter, report_data, tmp_path)
        assert doc is not None
        assert len(doc.paragraphs) > 0, "Document must contain paragraphs"


# ---------------------------------------------------------------------------
# 3-4: Title page and TOC
# ---------------------------------------------------------------------------


class TestTitlePageAndToc:
    """Verify title page content and Table of Contents section."""

    def test_title_page_contains_query_text(
        self, exporter, report_data, tmp_path,
    ):
        """The query/original_query must appear on the title page."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)
        full = _full_text(doc)
        query = report_data.get("original_query", report_data.get("query", ""))
        assert query in full, (
            f"Query text not found in document. "
            f"Expected: '{query[:80]}...'"
        )

    def test_table_of_contents_heading_present(
        self, exporter, report_data, tmp_path,
    ):
        """A 'Table of Contents' heading must exist in the document."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)
        full = _full_text(doc)
        assert "Table of Contents" in full, (
            "Missing 'Table of Contents' heading in document"
        )


# ---------------------------------------------------------------------------
# 5-6: Body sections and bibliography
# ---------------------------------------------------------------------------


class TestBodyAndBibliography:
    """Verify report body sections and bibliography rendering."""

    def test_report_sections_present(
        self, exporter, report_data, tmp_path,
    ):
        """Each structured section title must appear in the document."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)
        full = _full_text(doc)

        sections = report_data.get("sections", [])
        if sections:
            for section in sections:
                title = section.get("title", "")
                if title:
                    assert title in full, (
                        f"Section title '{title}' not found in document body"
                    )

    def test_section_count_matches(
        self, exporter, report_data, tmp_path,
    ):
        """The number of Heading 2 paragraphs must be >= number of structured sections."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)

        heading2_count = sum(
            1 for p in doc.paragraphs
            if p.style and p.style.name == "Heading 2"
        )

        sections = report_data.get("sections", [])
        if sections:
            # Structured sections render as Heading 2; final_report may add more
            assert heading2_count >= len(sections), (
                f"Found {heading2_count} Heading 2 paragraphs but expected "
                f">= {len(sections)} structured sections"
            )

    def test_bibliography_heading_and_entries(
        self, exporter, report_data, tmp_path,
    ):
        """Bibliography heading must exist with numbered reference entries."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)
        full = _full_text(doc)

        assert "Bibliography" in full, "Missing 'Bibliography' heading"

        bibliography = report_data.get("bibliography", [])
        if bibliography:
            # Check that at least the first entry title appears
            first_title = bibliography[0].get("title", "")
            if first_title:
                # Bibliography entries may be in paragraph text or table cells
                all_content = _all_text(doc)
                assert first_title in all_content, (
                    f"First bibliography entry '{first_title}' not found"
                )


# ---------------------------------------------------------------------------
# 7-8: Quality summary and audit certificate
# ---------------------------------------------------------------------------


class TestQualityAndAudit:
    """Verify quality summary metrics and audit certificate."""

    def test_quality_summary_with_faithfulness(
        self, exporter, report_data, tmp_path,
    ):
        """Quality Summary heading must exist with the faithfulness percentage."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)
        full = _full_text(doc)

        assert "Quality Summary" in full, "Missing 'Quality Summary' heading"

        # The faithfulness score appears in the quality summary table
        qm = report_data.get("quality_metrics") or {}
        faith = qm.get(
            "faithfulness_score",
            report_data.get("faithfulness_score", 0.0),
        )
        faith_pct = f"{faith * 100:.1f}%"
        all_content = _all_text(doc)
        assert faith_pct in all_content, (
            f"Faithfulness value '{faith_pct}' not found in Quality Summary"
        )

    def test_audit_certificate_with_vector_id_and_sha256(
        self, exporter, report_data, tmp_path,
    ):
        """Audit Certificate must contain the vector_id and a SHA-256 hash."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)
        # Use _all_text to capture table cell content (audit certificate
        # renders its fields inside a two-column table, not paragraphs)
        all_content = _all_text(doc)

        assert "Audit Certificate" in all_content, (
            "Missing 'Audit Certificate' heading"
        )

        vector_id = report_data.get("vector_id", "")
        assert vector_id in all_content, (
            f"Vector ID '{vector_id}' not found in Audit Certificate"
        )

        # Verify the SHA-256 hash matches the final_report content
        final_report = report_data.get("final_report", "")
        expected_hash = hashlib.sha256(
            final_report.encode("utf-8")
        ).hexdigest()
        assert expected_hash in all_content, (
            f"SHA-256 hash '{expected_hash[:16]}...' not found in "
            f"Audit Certificate table cells"
        )


# ---------------------------------------------------------------------------
# 9: Citation superscript rendering
# ---------------------------------------------------------------------------


class TestCitationRendering:
    """Verify that inline [N] citations are rendered as superscript runs."""

    def test_citations_rendered_as_superscript(
        self, exporter, report_data, tmp_path,
    ):
        """At least one citation run must have font.superscript == True."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)

        superscript_citations = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.superscript and "[" in run.text and "]" in run.text:
                    superscript_citations.append(run.text)

        assert len(superscript_citations) > 0, (
            "No superscript citation runs found. The exporter must render "
            "inline [N] references as superscript."
        )


# ---------------------------------------------------------------------------
# 10: Large report stress test
# ---------------------------------------------------------------------------


class TestLargeReport:
    """Verify the exporter handles large documents without error."""

    def test_large_report_exports_without_error(
        self, exporter, tmp_path,
    ):
        """A report with 15 sections must export cleanly."""
        large_data = {
            "vector_id": "V_LARGE_001",
            "original_query": "Comprehensive review of water treatment technologies",
            "status": "completed",
            "final_report": "",
            "sections": [],
            "bibliography": [],
            "quality_metrics": {
                "faithfulness_score": 0.92,
                "total_evidence": 500,
                "unique_sources": 80,
                "total_words": 25000,
                "total_sections": 15,
                "total_citations": 120,
                "coverage_score": 0.95,
                "coherence_score": 0.91,
                "gold_evidence": 60,
                "silver_evidence": 120,
            },
            "timestamps": {"created": "2024-06-01T08:00:00"},
            "iteration_count": 5,
            "convergence_reason": "max_iterations_reached",
            "quality_gate_result": "PASS",
            "faithfulness_score": 0.92,
            "evidence_count": 500,
        }

        # Generate 15 substantial sections
        for i in range(1, 16):
            section_content = (
                f"This section discusses aspect {i} of water treatment. "
                f"Research indicates that method {i} achieves significant "
                f"removal efficiency [{i}] across multiple contaminant classes. "
                f"A **meta-analysis** of {10 + i} studies confirms these "
                f"findings with *high* statistical confidence. "
            ) * 8  # ~600 words per section
            large_data["sections"].append({
                "title": f"Section {i}: Treatment Method Analysis",
                "content": section_content,
            })
            large_data["bibliography"].append({
                "citation_key": f"[{i}]",
                "title": f"Study on treatment method {i}",
                "url": f"https://example.com/study/{i}",
                "quality_tier": "GOLD" if i <= 5 else "SILVER",
            })

        # Build final_report from sections
        lines = []
        for sec in large_data["sections"]:
            lines.append(f"## {sec['title']}\n\n{sec['content']}\n")
        large_data["final_report"] = "\n".join(lines)

        output_path = tmp_path / "large_report.docx"
        result_path = exporter.export(large_data, output_path)

        assert result_path.exists()
        assert result_path.stat().st_size > 0

        # Verify all 15 sections rendered
        doc = DocxDocument(str(result_path))
        full = _full_text(doc)
        for i in range(1, 16):
            assert f"Section {i}:" in full, (
                f"Section {i} title missing from large report"
            )


# ---------------------------------------------------------------------------
# 11: Empty bibliography
# ---------------------------------------------------------------------------


class TestEdgeCases:
    """Edge cases that must not crash the exporter."""

    def test_empty_bibliography_exports_successfully(
        self, exporter, report_data, tmp_path,
    ):
        """A report with no bibliography entries must still produce valid .docx."""
        data = dict(report_data)
        data["bibliography"] = []
        data["vector_id"] = "V_EMPTY_BIB"

        output_path = tmp_path / "empty_bib.docx"
        result_path = exporter.export(data, output_path)
        assert result_path.exists()

        doc = DocxDocument(str(result_path))
        full = _full_text(doc)
        assert "Bibliography" in full
        assert "No bibliography entries available" in full

    def test_special_characters_in_query(self, exporter, tmp_path):
        """Special characters in the query must not crash the exporter."""
        data = {
            "vector_id": "V_SPECIAL_001",
            "original_query": (
                'What is the effect of "nano-TiO2" on <biofilm> '
                "formation at pH > 7 & temperature <= 25C?"
            ),
            "status": "completed",
            "final_report": (
                "# Results\n\n"
                'The effect of "nano-TiO2" at pH > 7 & temp <= 25C '
                "was studied [1].\n\n"
                "[1] Source with special chars: <brackets> & ampersand"
            ),
            "bibliography": [
                {
                    "citation_key": "[1]",
                    "title": 'Study of "nano-TiO2" <biofilm> & pH effects',
                    "url": "https://example.com/special?a=1&b=2",
                    "quality_tier": "SILVER",
                },
            ],
            "quality_metrics": {
                "faithfulness_score": 0.75,
                "total_evidence": 30,
                "unique_sources": 5,
                "total_words": 500,
                "total_sections": 1,
                "total_citations": 5,
                "coverage_score": 0.6,
                "coherence_score": 0.7,
                "gold_evidence": 3,
                "silver_evidence": 10,
            },
            "timestamps": {"created": "2024-03-01T12:00:00"},
            "iteration_count": 1,
            "convergence_reason": "synthesis_complete",
            "quality_gate_result": "PASS",
            "faithfulness_score": 0.75,
            "evidence_count": 30,
        }

        output_path = tmp_path / "special_chars.docx"
        result_path = exporter.export(data, output_path)
        assert result_path.exists()

        doc = DocxDocument(str(result_path))
        full = _full_text(doc)
        assert "nano-TiO2" in full

    def test_final_report_only_no_sections(self, exporter, tmp_path):
        """Export with final_report but empty sections list must succeed."""
        data = {
            "vector_id": "V_REPORT_ONLY",
            "original_query": "Test query for report-only path",
            "status": "completed",
            "final_report": (
                "# Summary\n\n"
                "This report was generated from final_report only.\n\n"
                "## Details\n\n"
                "No structured sections were provided [1].\n\n"
                "[1] Test source"
            ),
            "sections": [],
            "bibliography": [
                {
                    "citation_key": "[1]",
                    "title": "Test source",
                    "url": "https://example.com",
                    "quality_tier": "BRONZE",
                },
            ],
            "quality_metrics": {
                "faithfulness_score": 0.5,
                "total_evidence": 10,
                "unique_sources": 2,
                "total_words": 200,
                "total_sections": 1,
                "total_citations": 2,
                "coverage_score": 0.4,
                "coherence_score": 0.5,
                "gold_evidence": 1,
                "silver_evidence": 3,
            },
            "timestamps": {"created": "2024-02-20T09:00:00"},
            "iteration_count": 1,
            "convergence_reason": "synthesis_complete",
            "quality_gate_result": "PASS",
            "faithfulness_score": 0.5,
            "evidence_count": 10,
        }

        output_path = tmp_path / "report_only.docx"
        result_path = exporter.export(data, output_path)
        assert result_path.exists()

        doc = DocxDocument(str(result_path))
        full = _full_text(doc)
        assert "Summary" in full
        assert "final_report only" in full


# ---------------------------------------------------------------------------
# 12: File size constraint
# ---------------------------------------------------------------------------


class TestFileSizeConstraint:
    """Verify exported .docx stays within reasonable bounds."""

    def test_file_size_under_5mb(self, exporter, report_data, tmp_path):
        """The exported .docx must be smaller than 5 MB."""
        output_path = tmp_path / "size_check.docx"
        result_path = exporter.export(report_data, output_path)

        size_bytes = result_path.stat().st_size
        five_mb = 5 * 1024 * 1024
        assert size_bytes < five_mb, (
            f"File size {size_bytes / 1024 / 1024:.2f} MB exceeds 5 MB limit"
        )


# ---------------------------------------------------------------------------
# 13-15: Validation errors
# ---------------------------------------------------------------------------


class TestValidation:
    """Verify that invalid input triggers proper ValueError exceptions."""

    def test_missing_vector_id_raises_value_error(self, exporter, tmp_path):
        """Omitting 'vector_id' from report_data must raise ValueError."""
        bad_data = {
            "status": "completed",
            "final_report": "# Title\n\nSome content.",
        }
        output_path = tmp_path / "no_vector_id.docx"
        with pytest.raises(ValueError, match="vector_id"):
            exporter.export(bad_data, output_path)

    def test_missing_both_report_and_sections_raises_value_error(
        self, exporter, tmp_path,
    ):
        """Omitting both 'final_report' and 'sections' must raise ValueError."""
        bad_data = {
            "vector_id": "V_NO_BODY",
            "status": "completed",
        }
        output_path = tmp_path / "no_body.docx"
        with pytest.raises(ValueError, match="final_report.*sections|sections.*final_report"):
            exporter.export(bad_data, output_path)

    def test_missing_status_raises_value_error(self, exporter, tmp_path):
        """Omitting 'status' from report_data must raise ValueError."""
        bad_data = {
            "vector_id": "V_NO_STATUS",
            "final_report": "# Title\n\nSome content.",
        }
        output_path = tmp_path / "no_status.docx"
        with pytest.raises(ValueError, match="status"):
            exporter.export(bad_data, output_path)


# ---------------------------------------------------------------------------
# 16: Bold text rendering
# ---------------------------------------------------------------------------


class TestInlineFormatting:
    """Verify markdown inline formatting is translated to docx runs."""

    def test_bold_text_rendered(self, exporter, report_data, tmp_path):
        """Text wrapped in **double asterisks** must produce a bold run."""
        _, doc = _export_and_load(exporter, report_data, tmp_path)

        bold_runs = []
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.bold and run.text.strip():
                    bold_runs.append(run.text)

        assert len(bold_runs) > 0, (
            "No bold runs found. The exporter must render **text** as bold."
        )
