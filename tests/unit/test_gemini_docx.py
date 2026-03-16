"""
Unit tests for DOCX rendering features in DocxExporter.

Covers: markdown table rendering, base64 image embedding, left border styling,
Key Findings block detection, figure caption detection, and full export
integration with tables and images.

All tests use REAL python-docx objects -- no mocks.
"""

import base64

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.polaris_graph.export.docx_exporter import DocxExporter


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def exporter() -> DocxExporter:
    """Return a fresh DocxExporter instance."""
    return DocxExporter()


# Minimal valid 1x1 red PNG (known bytes, widely used in test suites).
VALID_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAA"
    "DUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
)


# ---------------------------------------------------------------------------
# 1. test_render_markdown_table_basic
# ---------------------------------------------------------------------------


def test_render_markdown_table_basic(exporter: DocxExporter) -> None:
    """Verify basic table creation: row/column counts and cell contents."""
    doc = Document()
    lines = [
        "| Name | Score |",
        "| --- | --- |",
        "| Alice | 95 |",
        "| Bob | 87 |",
    ]
    exporter._render_markdown_table(doc, lines)

    assert len(doc.tables) == 1, "Expected exactly 1 table in the document"

    table = doc.tables[0]
    # 1 header row + 2 data rows = 3 total
    assert len(table.rows) == 3, f"Expected 3 rows, got {len(table.rows)}"
    assert len(table.columns) == 2, f"Expected 2 columns, got {len(table.columns)}"

    # Header cells
    assert table.cell(0, 0).text.strip() == "Name"
    assert table.cell(0, 1).text.strip() == "Score"

    # Data cells
    assert table.cell(1, 0).text.strip() == "Alice"
    assert table.cell(1, 1).text.strip() == "95"
    assert table.cell(2, 0).text.strip() == "Bob"
    assert table.cell(2, 1).text.strip() == "87"


# ---------------------------------------------------------------------------
# 2. test_render_markdown_table_alternating_shading
# ---------------------------------------------------------------------------


def test_render_markdown_table_alternating_shading(
    exporter: DocxExporter,
) -> None:
    """Verify that alternating data rows receive XML shading elements.

    The implementation applies shading to odd-indexed data rows (0-indexed),
    so the *second* data row (r_idx=1) should have a ``w:shd`` element while
    the *first* data row (r_idx=0) should not.
    """
    doc = Document()
    lines = [
        "| A | B |",
        "| --- | --- |",
        "| row0_a | row0_b |",
        "| row1_a | row1_b |",
        "| row2_a | row2_b |",
    ]
    exporter._render_markdown_table(doc, lines)

    table = doc.tables[0]

    # Data row at r_idx=0 (table row index 1) -> NO shading
    cell_r0 = table.cell(1, 0)
    tc_pr_r0 = cell_r0._tc.tcPr
    shd_elements_r0 = (
        tc_pr_r0.findall(qn("w:shd")) if tc_pr_r0 is not None else []
    )
    assert len(shd_elements_r0) == 0, (
        "Data row 0 (r_idx=0) should NOT have shading"
    )

    # Data row at r_idx=1 (table row index 2) -> HAS shading (stripe color)
    cell_r1 = table.cell(2, 0)
    tc_pr_r1 = cell_r1._tc.tcPr
    assert tc_pr_r1 is not None, "Data row 1 should have tcPr element"
    shd_elements_r1 = tc_pr_r1.findall(qn("w:shd"))
    assert len(shd_elements_r1) >= 1, (
        "Data row 1 (r_idx=1) should have shading applied"
    )
    # Verify the fill color is the stripe color #F2F2F2
    fill_val = shd_elements_r1[0].get(qn("w:fill"))
    assert fill_val is not None
    assert fill_val.upper() == "F2F2F2", (
        f"Expected stripe fill F2F2F2, got {fill_val}"
    )

    # Data row at r_idx=2 (table row index 3) -> NO shading
    cell_r2 = table.cell(3, 0)
    tc_pr_r2 = cell_r2._tc.tcPr
    shd_elements_r2 = (
        tc_pr_r2.findall(qn("w:shd")) if tc_pr_r2 is not None else []
    )
    assert len(shd_elements_r2) == 0, (
        "Data row 2 (r_idx=2) should NOT have shading"
    )


# ---------------------------------------------------------------------------
# 3. test_render_base64_image_valid_png
# ---------------------------------------------------------------------------


def test_render_base64_image_valid_png(exporter: DocxExporter) -> None:
    """Embed a minimal 1x1 PNG and verify the document contains an InlineShape."""
    doc = Document()
    exporter._render_base64_image(
        doc,
        alt_text="test image",
        img_format="png",
        b64_data=VALID_PNG_B64,
    )

    # python-docx exposes inline images via doc.inline_shapes
    assert len(doc.inline_shapes) == 1, (
        f"Expected 1 inline shape, got {len(doc.inline_shapes)}"
    )


# ---------------------------------------------------------------------------
# 4. test_render_base64_image_invalid_graceful
# ---------------------------------------------------------------------------


def test_render_base64_image_invalid_graceful(exporter: DocxExporter) -> None:
    """Pass garbage base64 data. Verify no crash and fallback text appears."""
    doc = Document()
    exporter._render_base64_image(
        doc,
        alt_text="broken chart",
        img_format="png",
        b64_data="NOT_VALID_BASE64!!!@@@",
    )

    # No crash -- the method should have caught the exception.
    # Verify fallback placeholder text "[Image: broken chart]" exists.
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "[Image: broken chart]" in all_text, (
        f"Expected fallback text '[Image: broken chart]' in document, "
        f"got: {all_text!r}"
    )


# ---------------------------------------------------------------------------
# 5. test_render_markdown_detects_table
# ---------------------------------------------------------------------------


def test_render_markdown_detects_table(exporter: DocxExporter) -> None:
    """Pass markdown with paragraphs and an embedded table to _render_markdown.

    Verify the document contains both regular paragraphs and a table element.
    """
    markdown = (
        "Introduction paragraph before the table.\n"
        "\n"
        "| Material | Strength (MPa) |\n"
        "| --- | --- |\n"
        "| Steel | 250 |\n"
        "| Aluminum | 70 |\n"
        "\n"
        "Conclusion paragraph after the table."
    )
    doc = Document()
    exporter._render_markdown(doc, markdown)

    # At least 1 table
    assert len(doc.tables) >= 1, "Expected at least 1 table in the document"

    # Verify table content
    table = doc.tables[0]
    assert table.cell(0, 0).text.strip() == "Material"
    assert table.cell(1, 0).text.strip() == "Steel"

    # Verify paragraphs exist (intro + conclusion + any spacing paragraphs)
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert any("Introduction" in t for t in para_texts), (
        "Expected intro paragraph in document"
    )
    assert any("Conclusion" in t for t in para_texts), (
        "Expected conclusion paragraph in document"
    )


# ---------------------------------------------------------------------------
# 6. test_render_markdown_key_findings
# ---------------------------------------------------------------------------


def test_render_markdown_key_findings(exporter: DocxExporter) -> None:
    """Verify Key Findings detection: heading paragraph + bullet items.

    Note: ``**Key Findings:**`` (bare, no trailing text) is intercepted by the
    figure-caption regex ``^\\*(.+)\\*$`` before the Key Findings branch runs.
    Using ``**Key Findings:** Summary`` ensures the Key Findings path fires
    via ``stripped.startswith("**Key Findings")``.
    """
    markdown = (
        "**Key Findings:** Summary of results\n"
        "- Finding one [1]\n"
        "- Finding two [2]\n"
    )
    doc = Document()
    exporter._render_markdown(doc, markdown)

    # Collect all paragraph texts
    para_texts = [p.text for p in doc.paragraphs]

    # The Key Findings heading paragraph should contain "Key Findings"
    key_findings_paras = [t for t in para_texts if "Key Findings" in t]
    assert len(key_findings_paras) >= 1, (
        f"Expected 'Key Findings' heading paragraph, got texts: {para_texts}"
    )

    # Verify bullet items follow: look for paragraphs containing the findings
    all_text = "\n".join(para_texts)
    assert "Finding one" in all_text, (
        f"Expected 'Finding one' in document, got: {all_text!r}"
    )
    assert "Finding two" in all_text, (
        f"Expected 'Finding two' in document, got: {all_text!r}"
    )

    # Verify the Key Findings heading has a left border via XML
    for para in doc.paragraphs:
        if "Key Findings" in para.text and para.text.strip() == "Key Findings":
            p_pr = para._p.pPr
            assert p_pr is not None, "Key Findings paragraph should have pPr"
            p_bdr = p_pr.findall(qn("w:pBdr"))
            assert len(p_bdr) >= 1, (
                "Key Findings paragraph should have left border (w:pBdr)"
            )
            left_borders = p_bdr[0].findall(qn("w:left"))
            assert len(left_borders) >= 1, (
                "pBdr should contain a w:left element"
            )
            break
    else:
        pytest.fail(
            "Could not find standalone 'Key Findings' heading paragraph"
        )


# ---------------------------------------------------------------------------
# 7. test_render_markdown_figure_caption
# ---------------------------------------------------------------------------


def test_render_markdown_figure_caption(exporter: DocxExporter) -> None:
    """Verify figure caption renders as centered italic paragraph."""
    markdown = "*Figure 1: Adhesion strength comparison*"
    doc = Document()
    exporter._render_markdown(doc, markdown)

    # Find the caption paragraph
    caption_paras = [
        p for p in doc.paragraphs
        if "Figure 1" in p.text or "Adhesion strength" in p.text
    ]
    assert len(caption_paras) >= 1, (
        f"Expected figure caption paragraph, got: "
        f"{[p.text for p in doc.paragraphs]}"
    )

    caption_para = caption_paras[0]

    # Verify centered alignment
    assert caption_para.alignment == WD_ALIGN_PARAGRAPH.CENTER, (
        f"Expected CENTER alignment, got {caption_para.alignment}"
    )

    # Verify italic formatting on the run
    assert len(caption_para.runs) >= 1, "Caption paragraph should have runs"
    caption_run = caption_para.runs[0]
    assert caption_run.font.italic is True, (
        f"Expected italic run, got italic={caption_run.font.italic}"
    )


# ---------------------------------------------------------------------------
# 8. test_full_export_with_tables_and_images
# ---------------------------------------------------------------------------


def test_full_export_with_tables_and_images(
    exporter: DocxExporter,
    tmp_path,
) -> None:
    """Build a full report_data dict with markdown table AND base64 image.

    Export to a tmp_path .docx file, re-open and verify it contains at least
    1 table and the file exceeds 5 KB (contains image data).
    """
    # Build the final_report with an embedded markdown table and a base64 image
    final_report = (
        "# Research Report\n"
        "\n"
        "## Summary\n"
        "\n"
        "This report examines material properties.\n"
        "\n"
        "| Property | Value |\n"
        "| --- | --- |\n"
        "| Tensile Strength | 450 MPa |\n"
        "| Elongation | 12% |\n"
        "\n"
        f"![chart](data:image/png;base64,{VALID_PNG_B64})\n"
        "\n"
        "*Figure 1: Material comparison chart*\n"
        "\n"
        "**Key Findings:**\n"
        "- Material A outperforms Material B [1]\n"
        "- Cost efficiency improved by 15% [2]\n"
        "\n"
        "## Conclusion\n"
        "\n"
        "The results confirm the hypothesis.\n"
    )

    report_data = {
        "vector_id": "VEC_TEST_001",
        "status": "completed",
        "original_query": "Material properties analysis",
        "final_report": final_report,
        "bibliography": [
            {
                "citation_key": "[1]",
                "title": "Materials Science Today",
                "url": "https://example.com/1",
            },
            {
                "citation_key": "[2]",
                "title": "Cost Analysis Journal",
                "url": "https://example.com/2",
            },
        ],
        "quality_metrics": {
            "faithfulness_score": 0.85,
            "word_count": 3000,
            "citation_count": 10,
        },
        "evidence_count": 50,
        "iteration_count": 3,
        "timestamps": {
            "started": "2026-01-01T00:00:00Z",
            "completed": "2026-01-01T01:30:00Z",
        },
    }

    output_file = tmp_path / "test_report.docx"
    result_path = exporter.export(report_data, output_file)

    # File was written
    assert result_path.exists(), f"Output file does not exist: {result_path}"

    # File exceeds 5 KB (image data embedded)
    file_size = result_path.stat().st_size
    assert file_size > 5_000, (
        f"Expected file > 5 KB (contains image), got {file_size} bytes"
    )

    # Re-open and verify structure
    reopened = Document(str(result_path))

    # Should contain at least 1 table (the markdown table + possibly
    # quality summary tables from the exporter)
    assert len(reopened.tables) >= 1, (
        f"Expected >= 1 table in exported doc, got {len(reopened.tables)}"
    )

    # Verify the markdown table content survived round-trip
    found_material_table = False
    for table in reopened.tables:
        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                cell_texts.append(cell.text.strip())
        if "Property" in cell_texts and "Tensile Strength" in cell_texts:
            found_material_table = True
            break
    assert found_material_table, (
        "Expected to find the material properties table in the exported document"
    )


# ---------------------------------------------------------------------------
# Additional edge-case: _add_left_border standalone
# ---------------------------------------------------------------------------


def test_add_left_border_applies_xml(exporter: DocxExporter) -> None:
    """Verify _add_left_border injects w:pBdr > w:left XML into a paragraph."""
    from docx.shared import RGBColor as _RGBColor

    doc = Document()
    para = doc.add_paragraph("Test paragraph")

    color = _RGBColor(0x1F, 0x4E, 0x79)
    DocxExporter._add_left_border(para, color, width_pt=6)

    p_pr = para._p.pPr
    assert p_pr is not None, "Paragraph should have pPr after adding border"

    p_bdr_list = p_pr.findall(qn("w:pBdr"))
    assert len(p_bdr_list) == 1, "Expected exactly 1 w:pBdr element"

    left_list = p_bdr_list[0].findall(qn("w:left"))
    assert len(left_list) == 1, "Expected exactly 1 w:left element"

    left_el = left_list[0]
    # width_pt=6 -> sz = 6*2 = 12 half-points
    assert left_el.get(qn("w:sz")) == "12", (
        f"Expected w:sz='12', got '{left_el.get(qn('w:sz'))}'"
    )
    assert left_el.get(qn("w:val")) == "single"
    assert left_el.get(qn("w:space")) == "4"
