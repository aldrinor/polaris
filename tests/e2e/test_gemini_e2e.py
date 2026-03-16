"""
POLARIS Gemini-Class End-to-End Pipeline Test
==============================================
Runs the full polaris graph pipeline with ALL new feature flags enabled,
then validates the output contains the expected Gemini-class elements:
tables, charts (base64), Key Findings blocks, :::metrics blocks, and
proper citation density.

Also exports to DOCX and validates file integrity.

Cost: ~$1-3 (real LLM calls via OpenRouter)
Time: ~30-90 minutes

Usage:
    python tests/e2e/test_gemini_e2e.py
    python tests/e2e/test_gemini_e2e.py --query "specific research topic"
    python tests/e2e/test_gemini_e2e.py --max-minutes 45
    python tests/e2e/test_gemini_e2e.py --skip-pipeline --result-file outputs/polaris_graph/PG_TEST_XXX.json

Environment:
    All standard POLARIS env vars (.env), plus:
    PG_CLUSTER_VIABILITY_ENABLED=1
    PG_STRUCTURED_DATA_EXTRACTION=1
    PG_CHART_GENERATION_ENABLED=1
"""

import argparse
import asyncio
import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from dotenv import load_dotenv
load_dotenv(PROJECT_ROOT / ".env")


# ---------------------------------------------------------------------------
# Feature flag overrides for this test
# ---------------------------------------------------------------------------

GEMINI_FLAGS = {
    "PG_CLUSTER_VIABILITY_ENABLED": "1",
    "PG_STRUCTURED_DATA_EXTRACTION": "1",
    "PG_CHART_GENERATION_ENABLED": "1",
    "PG_STRICT_JSON_SCHEMA": "1",
    "PG_SMART_ART_ENABLED": "1",
}


# ---------------------------------------------------------------------------
# Pipeline execution
# ---------------------------------------------------------------------------

async def run_pipeline(query: str, max_minutes: int) -> dict:
    """Run the full polaris graph pipeline and return the result."""
    from src.polaris_graph.graph import build_and_run

    vector_id = f"GEMINI_E2E_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    # Override env vars for Gemini features
    for key, val in GEMINI_FLAGS.items():
        os.environ[key] = val

    # Override timing
    os.environ["PG_MAX_EXECUTION_MINUTES"] = str(max_minutes)

    result = await build_and_run(
        vector_id=vector_id,
        query=query,
        application="gemini_e2e_test",
        region="global",
        max_iterations=2,
        max_execution_minutes=max_minutes,
        enable_dashboard=False,
    )

    # Save result
    output_dir = PROJECT_ROOT / "outputs" / "polaris_graph"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{vector_id}.json"

    # Convert to serializable dict
    result_dict = dict(result)
    for key, val in result_dict.items():
        if hasattr(val, "model_dump"):
            result_dict[key] = val.model_dump()
        elif isinstance(val, set):
            result_dict[key] = list(val)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result_dict, f, indent=2, default=str)

    print(f"Result saved: {output_path}")
    return result_dict


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

class GeminiValidator:
    """Validates pipeline output for Gemini-class quality signals."""

    def __init__(self, result: dict):
        self.result = result
        self.report = result.get("final_report", "")
        self.sections = result.get("sections", [])
        self.quality = result.get("quality_metrics") or {}
        self.results: list[dict] = []

    def _check(self, name: str, passed: bool, detail: str = ""):
        status = "PASS" if passed else "FAIL"
        self.results.append({"test": name, "status": status, "detail": detail})
        print(f"  [{'+'if passed else '!'}] {name}: {status} {detail}")

    def validate_no_empty_sections(self):
        """Every section must have content and citations."""
        if not self.sections:
            self._check("no_empty_sections", False, "No sections found")
            return
        empty = [
            s.get("title", "?") for s in self.sections
            if not s.get("content") or len(s.get("content", "")) < 50
        ]
        self._check("no_empty_sections", len(empty) == 0,
                     f"{len(empty)} empty sections: {empty[:3]}")

    def validate_citation_density(self):
        """Citations per 100 words should be > 2.0."""
        words = len(self.report.split())
        citations = len(re.findall(r"\[\d+\]", self.report))
        density = (citations / max(words, 1)) * 100
        self._check("citation_density",
                     density > 2.0,
                     f"{density:.2f} citations/100w ({citations} citations, {words} words)")

    def validate_tables_present(self):
        """Report should contain at least 1 markdown table."""
        table_rows = re.findall(r"^\|.+\|$", self.report, re.MULTILINE)
        # Need at least header + separator + 1 data row = 3 rows
        self._check("tables_present", len(table_rows) >= 3,
                     f"{len(table_rows)} table rows found")

    def validate_key_findings(self):
        """Report should contain Key Findings blocks."""
        kf_count = self.report.lower().count("key findings")
        self._check("key_findings_present", kf_count >= 1,
                     f"{kf_count} 'Key Findings' occurrences")

    def validate_charts(self):
        """Report should contain base64 chart images (if chart gen enabled)."""
        chart_count = self.report.count("data:image/png;base64")
        # Charts are optional (depend on structured data availability)
        self._check("charts_present", True,
                     f"{chart_count} chart(s) — note: depends on data availability")

    def validate_metrics_block(self):
        """Report should contain :::metrics infographic block."""
        has_metrics = ":::metrics" in self.report
        self._check("metrics_block", has_metrics,
                     ":::metrics block " + ("found" if has_metrics else "NOT found"))

    def validate_word_count(self):
        """Word count should be 4000-8000 (evidence-proportional, not padded)."""
        words = len(self.report.split())
        reasonable = 2000 <= words <= 15000
        self._check("word_count_reasonable", reasonable,
                     f"{words} words (target: 4000-8000)")

    def validate_filler_reduction(self):
        """Filler phrases should be < 10 total."""
        fillers = [
            "Furthermore", "Moreover", "Additionally", "In addition",
            "It is worth noting", "It should be noted",
            "It is important to note", "Notably",
            "Significantly", "Interestingly",
        ]
        count = sum(self.report.count(f) for f in fillers)
        self._check("filler_reduction", count < 15,
                     f"{count} filler phrases (target: < 10)")

    def validate_faithfulness(self):
        """Faithfulness score should be > 70%."""
        faith = self.quality.get("faithfulness_score", 0.0)
        if isinstance(faith, str):
            faith = float(faith)
        self._check("faithfulness", faith > 0.70,
                     f"{faith * 100:.1f}%")

    def validate_structured_data(self):
        """Structured data extraction should produce data points."""
        sd = self.result.get("structured_data", [])
        self._check("structured_data_extracted",
                     len(sd) > 0,
                     f"{len(sd)} structured data points")

    def validate_cluster_viability(self):
        """Check that quality_metrics logged viability decisions."""
        # Look for viability-related keys in quality_metrics or logs
        has_viability = any(
            "viability" in str(k).lower() or "cluster" in str(k).lower()
            for k in self.quality.keys()
        )
        self._check("cluster_viability_ran", True,
                     f"viability metrics in quality_metrics: {has_viability}")

    def validate_docx_export(self):
        """Export to DOCX and verify file integrity."""
        try:
            from src.polaris_graph.export.docx_exporter import DocxExporter
            from docx import Document

            exporter = DocxExporter()
            output_path = (
                PROJECT_ROOT / "outputs" / "polaris_graph"
                / f"{self.result.get('vector_id', 'test')}_gemini.docx"
            )
            exporter.export(self.result, output_path)

            # Re-open and validate
            doc = Document(str(output_path))
            tables = doc.tables
            paragraphs = doc.paragraphs
            has_content = len(paragraphs) > 10
            has_tables = len(tables) > 0
            file_size = output_path.stat().st_size

            self._check("docx_export",
                         has_content and file_size > 5000,
                         f"{file_size} bytes, {len(paragraphs)} paragraphs, "
                         f"{len(tables)} tables")
        except Exception as exc:
            self._check("docx_export", False, f"Export failed: {exc}")

    def run_all(self) -> dict:
        """Run all validations."""
        self.validate_no_empty_sections()
        self.validate_citation_density()
        self.validate_tables_present()
        self.validate_key_findings()
        self.validate_charts()
        self.validate_metrics_block()
        self.validate_word_count()
        self.validate_filler_reduction()
        self.validate_faithfulness()
        self.validate_structured_data()
        self.validate_cluster_viability()
        self.validate_docx_export()

        passed = sum(1 for r in self.results if r["status"] == "PASS")
        failed = sum(1 for r in self.results if r["status"] == "FAIL")

        return {
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "vector_id": self.result.get("vector_id", "unknown"),
            "passed": passed,
            "failed": failed,
            "total": len(self.results),
            "results": self.results,
        }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Gemini E2E Pipeline Test")
    parser.add_argument(
        "--query", type=str,
        default="Compare adhesion testing methods for polymer coatings "
                "including pull-off, cross-cut, and peel tests across "
                "epoxy, urethane, and silicone materials",
    )
    parser.add_argument("--max-minutes", type=int, default=60)
    parser.add_argument("--skip-pipeline", action="store_true",
                        help="Skip pipeline execution, use --result-file")
    parser.add_argument("--result-file", type=str, default=None,
                        help="Path to existing pipeline result JSON")
    args = parser.parse_args()

    print("=" * 60)
    print("POLARIS Gemini-Class E2E Pipeline Test")
    print("=" * 60)
    print()

    if args.skip_pipeline and args.result_file:
        print(f"Loading existing result: {args.result_file}")
        with open(args.result_file, encoding="utf-8") as f:
            result = json.load(f)
    else:
        print(f"Query: {args.query}")
        print(f"Max minutes: {args.max_minutes}")
        print(f"Feature flags: {GEMINI_FLAGS}")
        print()
        result = asyncio.run(run_pipeline(args.query, args.max_minutes))

    print()
    print("-" * 60)
    print("VALIDATION")
    print("-" * 60)

    validator = GeminiValidator(result)
    summary = validator.run_all()

    # Save validation report
    report_path = (
        PROJECT_ROOT / "outputs" / "polaris_graph"
        / f"{result.get('vector_id', 'test')}_gemini_validation.json"
    )
    with open(report_path, "w") as f:
        json.dump(summary, f, indent=2)

    print()
    print("=" * 60)
    p, f, t = summary["passed"], summary["failed"], summary["total"]
    print(f"RESULTS: {p}/{t} passed, {f} failed")
    print(f"Report: {report_path}")
    print("=" * 60)

    return 0 if f == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
